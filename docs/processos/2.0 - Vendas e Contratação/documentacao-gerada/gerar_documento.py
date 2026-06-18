"""
Gerador de documentação para o processo 2.0 - Vendas e Contratação
Versão: 1.0
Executar: python3 gerar_documento.py
"""

import sys
from pathlib import Path

# ── nexforce-brand helpers ──────────────────────────────────────────────────
SKILL_DIR = Path("/home/hugo-zanni/Nexforce/Projects/handover-sienge/.claude/skills/nexforce-brand")
sys.path.insert(0, str(SKILL_DIR / "references"))

import python_docx_helpers as _brand
_brand.nexforce_assets = lambda: SKILL_DIR / "assets"

from python_docx_helpers import (
    apply_nexforce_theme,
    insert_brand_header,
    insert_brand_footer,
    create_brand_table,
    add_callout,
    NEXFORCE_COLORS,
    logo_path,
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "2.0 - Vendas e Contratação.docx"


def _rgb(key: str) -> RGBColor:
    return RGBColor.from_string(NEXFORCE_COLORS[key])


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Lato"
    if level == 1:
        run.font.color.rgb = _rgb("primary_black")
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = _rgb("navy")
        run.font.size = Pt(13)
        run.font.bold = True
    else:
        run.font.color.rgb = _rgb("dark_gray")
        run.font.size = Pt(11)
        run.font.bold = True
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color_key="dark_gray", indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Lato"
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color_key)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    for run in p.runs:
        run.font.name = "Lato"
        run.font.size = Pt(10.5)
        run.font.color.rgb = _rgb("dark_gray")
    return p


def add_workflow_block(doc, nome, id_wf, objetivo, objeto, disparo, acoes, dependencias=None, status="Ativo"):
    p = doc.add_paragraph()
    run = p.add_run(f"Workflow: {nome}")
    run.bold = True
    run.font.name = "Lato"
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb("navy")

    rows_data = [
        ("ID", id_wf),
        ("Objetivo", objetivo),
        ("Status", status),
        ("Objeto HubSpot", objeto),
        ("Disparo", disparo),
        ("Ações principais", acoes),
    ]
    if dependencias:
        rows_data.append(("Dependências", dependencias))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        cell_l = table.rows[i].cells[0]
        cell_l.text = label
        _set_cell_bg(cell_l, NEXFORCE_COLORS["callout_info_bg"])
        for para in cell_l.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Lato"
                run.font.size = Pt(9.5)
                run.font.color.rgb = _rgb("navy")

        cell_v = table.rows[i].cells[1]
        cell_v.text = value
        bg = NEXFORCE_COLORS["white"] if i % 2 == 0 else NEXFORCE_COLORS["near_white"]
        _set_cell_bg(cell_v, bg)
        for para in cell_v.paragraphs:
            for run in para.runs:
                run.font.name = "Lato"
                run.font.size = Pt(9.5)
                run.font.color.rgb = _rgb("dark_gray")

    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(13)
    doc.add_paragraph()


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
apply_nexforce_theme(doc)

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

insert_brand_header(doc, title="2.0 — Vendas e Contratação", period="Junho 2026")
insert_brand_footer(doc, confidentiality="Uso interno  ·  Nexforce Services  ·  v1.0")

# ════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Cm(8)
r = p_sub.add_run("Documentação de Processos — HubSpot RaaS")
r.font.name = "Lato"
r.font.size = Pt(12)
r.font.color.rgb = _rgb("mid_gray")

doc.add_paragraph()

p_title = doc.add_heading("2.0 — Vendas e Contratação", level=0)
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in p_title.runs:
    run.font.name = "Lato"
    run.font.color.rgb = _rgb("primary_black")
    run.font.size = Pt(26)

doc.add_paragraph()

meta_table = create_brand_table(
    doc,
    headers=["Campo", "Valor"],
    rows=[
        ["Cliente", "Sienge"],
        ["Responsável pelo processo", "Vinícius Vieira Braz"],
        ["Empresa responsável", "Nexforce Services"],
        ["Data de documentação", "16 de junho de 2026"],
        ["Versão", "1.0"],
        ["Status do processo", "Estabilização pós Go Live"],
    ],
)
set_col_widths(meta_table, [6, 10])

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SUMÁRIO (manual)
# ════════════════════════════════════════════════════════
add_heading(doc, "Sumário", level=1)
for item in [
    "1. Objetivo",
    "2. Contexto de Negócio",
    "3. Visão Funcional",
    "   3.1 Fluxo Operacional",
    "   3.2 Regras de Negócio",
    "   3.3 Critério de Validação",
    "4. Visão Técnica",
    "   4.1 Objetos HubSpot Envolvidos",
    "   4.2 Propriedades Críticas",
    "   4.3 Workflows Envolvidos",
    "   4.4 Cards e Customizações",
    "5. Riscos e Dependências",
    "6. Materiais de Apoio",
    "7. Pontos a Validar",
]:
    add_para(doc, item, size=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. OBJETIVO
# ════════════════════════════════════════════════════════
add_heading(doc, "1. Objetivo", level=1)
add_para(doc,
    "Conduzir o processo comercial completo de vendas no HubSpot — desde a criação do negócio até o "
    "fechamento (ganho ou perda) — garantindo que cada oportunidade percorra o pipeline correto conforme "
    "seu tipo (Aquisição, Expansão ou Retração), receba os responsáveis adequados via round robin, tenha "
    "seus contratos aprovados, assinados e registrados, e seja entregue ao time de CS e BackOffice com "
    "todas as informações necessárias para operação."
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 2. CONTEXTO DE NEGÓCIO
# ════════════════════════════════════════════════════════
add_heading(doc, "2. Contexto de Negócio", level=1)

add_para(doc,
    "O processo 2.0 é o núcleo da operação comercial da Sienge no HubSpot. Ele cobre três tipos de "
    "movimentação de negócio:"
)
for item in [
    "Aquisição: venda nova para cliente que ainda não é Sienge — gera pipeline de Aquisição (ID 795178644).",
    "Expansão: ampliação de contrato de cliente existente — gera pipeline de Expansão (ID 839459507).",
    "Retração: redução ou cancelamento de contrato de cliente existente — gera pipeline de Retração (ID 842042121).",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc,
    "Cada tipo tem seu próprio pipeline de deal no HubSpot, com etapas, SLAs, workflows de movimentação "
    "automática e fluxos de aprovação distintos. O processo 2.0 é a entrada para o processo 2.1 (geração "
    "de contratos e portfólio), para o processo 3.0 (aprovações) e para o processo 5.0 (minutas). "
    "O não cumprimento de um step neste processo bloqueia todos os processos subsequentes."
)
doc.add_paragraph()
add_para(doc,
    "O processo inclui as modalidades SaaS (Sienge Plataforma na nuvem) e LU com DC (Licença de Uso com "
    "Data Center, modelo legado), com regras distintas para aprovação, movimentação e geração de orçamento. "
    "A originação de Expansão e Retração pode vir do time de vendas (Inside Sales, BackOffice) ou do "
    "time de CS via playbook de cancelamento/redução preenchido no ticket de engajamento."
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 3. VISÃO FUNCIONAL
# ════════════════════════════════════════════════════════
add_heading(doc, "3. Visão Funcional", level=1)

# 3.1 Fluxo Operacional
add_heading(doc, "3.1 Fluxo Operacional", level=2)

add_para(doc, "Aquisição e Expansão", bold=True, color_key="navy")
add_para(doc, "O fluxo comercial de Aquisição e Expansão segue a sequência abaixo:")
doc.add_paragraph()

steps_aq = [
    ("1", "Lead ou Oportunidade é criada",
     "Para Aquisição: um Lead gerado via RD Station, formulário de entrada ou prospecção ativa é "
     "convertido em Deal de Aquisição automaticamente (ou criado manualmente). "
     "Para Expansão: o CS preenche o playbook ou o vendedor cria o Deal diretamente. "
     "Propriedades obrigatórias no Deal: origem, canal_responsavel (Canal Oficial), tipo_de_orcamento, "
     "modalidade_de_aquisicao."),
    ("2", "Round Robin — atribuição de equipe",
     "Assim que o Deal tem origem e canal_responsavel preenchidos, o workflow de Round Robin avalia a "
     "origem do Deal (Inside Sales, BackOffice ou Canal) e distribui o Deal entre os membros da equipe "
     "correspondente de forma rotativa. Os campos de equipe (equipe_vendas__inside_sales__vendedor, "
     "equipe_canal__vendedor, etc.) são preenchidos automaticamente. Workflow IDs: 1784845853 e 1798412036."),
    ("3", "Canal Oficial associado ao objeto Canal",
     "Após o preenchimento de canal_responsavel, o workflow de Associação (ID 1793578765) copia o "
     "canal_legado para o Deal e cria a associação entre o Deal e o objeto Canal (2-54226027) usando "
     "matching por canal_legado_oficiais__legado. "
     "O canal_legado é copiado da Empresa associada ao Deal, não do canal_responsavel."),
    ("4", "Progressão das etapas do pipeline",
     "O vendedor avança o Deal pelas etapas: Oportunidade Identificada → Oportunidade Qualificada → "
     "Agendamento → Solução → Negociação → Contrato. "
     "Na transição de Agendamento para Solução, o vendedor deve preencher Próximos Passos e "
     "Data dos Próximos Passos — automaticamente é criada uma tarefa com lembrete para essa data. "
     "Workflows de SLA notificam o responsável quando o Deal permanece além do prazo em cada etapa."),
    ("5", "Preenchimento do tipo e modalidade de orçamento",
     "Ao chegar em Negociação, o vendedor preenche tipo_de_orcamento e modalidade_de_aquisicao. "
     "Esses campos disparam o workflow de movimentação (ID 1753092609) que move o Deal para a "
     "etapa de Contrato correspondente (SaaS ou LU com DC)."),
    ("6", "Envio para aprovação e assinatura",
     "Na etapa de Contrato, o vendedor envia o orçamento (Quote) para aprovação. O processo de "
     "aprovação é tratado no processo 3.0. Aprovado o contrato, ele é enviado para assinatura digital "
     "via HubSpot. Todos os contratos devem incluir Maria Madalena e Giovani Amaral como signatários "
     "obrigatórios — caso ausentes, o contrato é reprovado automaticamente."),
    ("7", "Alerta de assinatura do cliente",
     "Quando o cliente (signatários externos) assina, um workflow dispara alerta para o time "
     "comercial e financeiro mesmo que Maria Madalena e Giovani Amaral ainda não tenham assinado. "
     "Isso permite iniciar o processo de liberação do produto sem aguardar os signatários internos."),
    ("8", "Ganho manual",
     "Quando todos os signatários do orçamento assinam e o campo gatilho_para_ganho_automatico = Sim, "
     "o workflow seta negocio_enviado_para_ganho_manual = Sim, preenche data_assinatura_contrato_pai "
     "e notifica o financeiro para dar o ganho manual no Deal. O Deal é então movido para Vendido "
     "(ou Upsell) manualmente pelo time financeiro. O ganho dispara o processo 2.1."),
    ("9", "Integração Freshdesk",
     "Ao fechar o Deal como ganho (Vendido), o campo freshdesk_data_do_contrato é preenchido "
     "automaticamente, disparando a integração com o Freshdesk para criação do cliente no sistema de CS."),
]

for num, titulo, descricao in steps_aq:
    p = doc.add_paragraph()
    r_num = p.add_run(f"Passo {num}: ")
    r_num.bold = True
    r_num.font.name = "Lato"
    r_num.font.color.rgb = _rgb("navy")
    r_num.font.size = Pt(11)
    r_titulo = p.add_run(titulo)
    r_titulo.bold = True
    r_titulo.font.name = "Lato"
    r_titulo.font.size = Pt(11)
    r_titulo.font.color.rgb = _rgb("primary_black")
    add_para(doc, descricao, size=10.5, indent=True)
    doc.add_paragraph()

add_para(doc, "Retração", bold=True, color_key="navy")
add_para(doc, "O fluxo de Retração (cancelamento ou redução) é originado pelo CS via playbook:")
doc.add_paragraph()

steps_ret = [
    ("1", "Playbook de Cancelamento e Redução preenchido",
     "O CS preenche o playbook 'Cancelamento ou Redução' no ticket de engajamento do cliente. "
     "O tipo de solicitação define o subtipo de retração (cancelamento, redução de módulo, "
     "redução temporária, etc.). Para tipo = Inadimplência, ver Ponto a Validar PV-01."),
    ("2", "Criação automática do Deal de Retração e cópia de canal",
     "O preenchimento do playbook dispara a criação automática de um Deal no pipeline de Retração "
     "(ID 842042121). As propriedades do playbook são copiadas do ticket de engajamento para o Deal. "
     "O canal_responsavel é copiado do ticket de engajamento para o Deal de retração."),
    ("3", "Geração automática do orçamento",
     "Para reduções temporárias e de conectores, o orçamento é gerado automaticamente. "
     "Após a geração, o Deal avança automaticamente para 'Orçamento Enviado'. "
     "O orçamento de retração é expirado automaticamente caso o Deal seja descartado."),
    ("4", "SLA de Retração",
     "Os SLAs do pipeline de Retração são:\n"
     "  12 dias corridos entre: data de início do aviso prévio e envio do contrato para aprovação do time de contratos (alerta: CS).\n"
     "  6 dias entre: recebimento da aprovação do time de contratos e aprovação do financeiro (alerta: CS e Financeiro).\n"
     "  6 dias entre: aprovação do financeiro e assinatura do cliente (alerta: CS).\n"
     "  5 dias após assinatura do cliente para o financeiro realizar os trâmites.\n"
     "  1 dia para registrar o ganho e mover para 'Retração Realizada'."),
    ("5", "Travas de passagem de etapa",
     "Campos booleanos (autorizar_passagem_para_contrato_assinado, autorizar_passagem_para_vendido) "
     "impedem que o CS mova o Deal de retração para etapas indevidas. "
     "O campo autorizar_passagem deve ser preenchido como 'Sim' pelo aprovador responsável para "
     "habilitar a movimentação."),
    ("6", "Data fim do aviso prévio",
     "O campo data_fim_do_aviso_previo pode ser preenchido manualmente (campo aberto) ou calculado. "
     "A referência base é 30 dias a partir da data de início do aviso prévio, mas o prazo real pode "
     "diferir por contrato. O campo é editável para permitir ajustes pontuais."),
]

for num, titulo, descricao in steps_ret:
    p = doc.add_paragraph()
    r_num = p.add_run(f"Passo {num}: ")
    r_num.bold = True
    r_num.font.name = "Lato"
    r_num.font.color.rgb = _rgb("navy")
    r_num.font.size = Pt(11)
    r_titulo = p.add_run(titulo)
    r_titulo.bold = True
    r_titulo.font.name = "Lato"
    r_titulo.font.size = Pt(11)
    r_titulo.font.color.rgb = _rgb("primary_black")
    add_para(doc, descricao, size=10.5, indent=True)
    doc.add_paragraph()

# 3.2 Regras de Negócio
add_heading(doc, "3.2 Regras de Negócio", level=2)

regras = [
    ("Regra 1 — ID Sienge: sem reaproveitamento entre produtos",
     "O campo id_sienge (ou id_do_cliente) representa o ID do produto específico contratado. "
     "Não se reutiliza o ID de um produto para outro. "
     "Exemplos:\n"
     "  Cliente com ID 3456 no Gestor Obras que cancela e contrata Sienge Plataforma → novo ID para Plataforma.\n"
     "  Cliente com ID 9567 na Sienge e ID 2345 na Prevision → IDs diferentes por produto, ambos mantidos.\n"
     "  Cliente que foi Plataforma, cancelou e volta para Plataforma → reaproveita o ID Plataforma existente.\n"
     "A propriedade deve espelhar o ID do produto que está ativo."),
    ("Regra 2 — Signatários obrigatórios em todos os contratos",
     "Todos os contratos enviados para assinatura via HubSpot devem incluir Maria Madalena e "
     "Giovani Amaral nos signatários. Se um contrato for enviado sem um deles, um workflow verifica "
     "a ausência e reprova o contrato automaticamente, notificando o solicitante com o motivo da reprovação."),
    ("Regra 3 — Alerta de assinatura do cliente antes dos signatários internos",
     "O financeiro e o time comercial consideram o contrato 'vendido' quando o cliente (signatários "
     "externos) assina, independentemente de Maria Madalena e Giovani Amaral terem assinado. "
     "Para refletir isso operacionalmente, quando todos os signatários externos assinam, um workflow "
     "notifica o time imediatamente, mesmo que os signatários internos ainda estejam pendentes."),
    ("Regra 4 — Canal Oficial e roteamento para aprovadores",
     "O canal que determina o aprovador de contratos em Aquisição é canal_implantacao (não canal_responsavel). "
     "Para Expansão, o canal_responsavel é o campo de roteamento. "
     "O campo canal_legado_oficiais__legado é copiado da Empresa associada ao Deal, nunca do canal_responsavel. "
     "Se o canal_responsavel for alterado por fluxo automático incorretamente, o contrato pode cair "
     "para o aprovador errado (risco documentado no ticket 86ba43rfc)."),
    ("Regra 5 — Round Robin por origem e canal",
     "A distribuição de Deals do pipeline de Aquisição para os vendedores segue a lógica:\n"
     "  Origem = 'Inside Sales' → distribui para a equipe de Inside Sales.\n"
     "  Origem = 'BackOffice' → distribui para a equipe de BackOffice.\n"
     "  Origem = 'Canal' (parceiro) → distribui para a equipe do Canal.\n"
     "O round robin para Expansão segue lógica similar com equipe_expansao__inside_sales__vendedor."),
    ("Regra 6 — LT Diário: conversão texto→número",
     "O campo LT Diário existe em duas versões no Deal: uma de texto e uma numérica. "
     "Um workflow converte automaticamente o valor preenchido no campo de texto para o campo numérico. "
     "Essa conversão é necessária para uso em cálculos e relatórios."),
    ("Regra 7 — Próximos Passos gera tarefa automática",
     "Quando o vendedor preenche Próximos Passos (proximos_passos) e Data dos Próximos Passos "
     "(data_da_proxima_etapa) no Deal, um workflow cria automaticamente uma tarefa no HubSpot "
     "com lembrete para a data informada. "
     "O campo Próximos Passos não pode ser apagado por nenhum fluxo automático."),
    ("Regra 8 — Ganho manual pelo financeiro",
     "O HubSpot não dá ganho automático ao Deal quando o contrato é assinado. "
     "O fluxo é: (1) contrato assinado → workflow notifica financeiro; (2) financeiro acessa o Deal "
     "e move manualmente para Vendido ou Upsell. "
     "O campo negocio_enviado_para_ganho_manual = Sim é o sinalizador de que o workflow já notificou "
     "o financeiro e o Deal aguarda a ação manual."),
    ("Regra 9 — Desconto Temporário e Desconto Escalonado",
     "Deals que envolvem desconto escalonado têm um responsável de aprovação específico "
     "(responsavel_pela_aprovacao_do_desconto_escalonado). O fluxo de aprovação de desconto é "
     "tratado como parte do processo 3.0 (Aprovações). "
     "O tipo 'Desconto Temporário' gera orçamento automaticamente para deals de retração com esse subtipo."),
    ("Regra 10 — Ticket de Retração por Inadimplência",
     "Quando o CS preenche o playbook de Cancelamento e Redução com tipo = Inadimplência, o fluxo "
     "deve criar um ticket de retração no pipeline. "
     "Esse caso está documentado como blocked (ticket 86bacyj4b) — não confirmado como funcionando."),
    ("Regra 11 — Troca de Canal",
     "Quando o CS cria um chamado de troca de canal via playbook, o sistema deve criar um ticket "
     "de troca de canal e enviar alerta para o time de contratos. "
     "Esse fluxo foi corrigido durante a estabilização (ticket 86ba8b0vy)."),
]

for titulo, texto in regras:
    add_para(doc, titulo, bold=True, color_key="primary_black")
    add_para(doc, texto, size=10.5, indent=True)
    doc.add_paragraph()

# 3.3 Critério de Validação
add_heading(doc, "3.3 Critério de Validação", level=2)
add_para(doc, "O processo foi executado corretamente quando:")
for item in [
    "O Deal existe no pipeline correto (Aquisição, Expansão ou Retração) com tipo_de_orcamento e modalidade_de_aquisicao preenchidos.",
    "O campo canal_responsavel (Canal Oficial) está preenchido e o Deal está associado ao objeto Canal correspondente.",
    "O canal_legado_oficiais__legado está preenchido com o valor da Empresa (não do canal_responsavel).",
    "O vendedor ou equipe responsável foi atribuído via round robin (equipe_* preenchida).",
    "O Deal avançou pelas etapas do pipeline sem ficar travado além do SLA configurado.",
    "O contrato enviado para assinatura inclui Maria Madalena e Giovani Amaral como signatários.",
    "Quando o cliente assina, o time comercial/financeiro recebe a notificação antes da assinatura dos signatários internos.",
    "O campo negocio_enviado_para_ganho_manual = Sim indica que o financeiro foi notificado.",
    "O Deal ganho disparou o processo 2.1 (criação de Contrato, GC e Portfólio).",
    "Para Retração: o canal_responsavel foi copiado corretamente do ticket de engajamento para o Deal.",
    "Para Retração: o SLA está sendo monitorado com alertas por etapa.",
]:
    add_bullet(doc, item)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 4. VISÃO TÉCNICA
# ════════════════════════════════════════════════════════
add_heading(doc, "4. Visão Técnica", level=1)

# 4.1 Objetos HubSpot
add_heading(doc, "4.1 Objetos HubSpot Envolvidos", level=2)
t41 = create_brand_table(doc,
    ["Objeto", "objectTypeId", "Papel no processo"],
    [
        ["Deal (Negócio)", "0-3",
         "Objeto central do processo 2.0. Existe em três pipelines: Aquisição (795178644), "
         "Expansão (839459507) e Retração (842042121). Dispara o processo 2.1 ao ser ganho."],
        ["Canal", "2-54226027",
         "Objeto customizado. Associado ao Deal via canal_legado_oficiais__legado. "
         "Define o aprovador de contratos e a equipe responsável."],
        ["Empresa (Company)", "0-2",
         "Origem do canal_legado copiado para o Deal. Associada ao Deal e ao processo 2.1."],
        ["Contato (Contact)", "0-1",
         "Associado ao Deal. Pode ser o contato de assinatura do contrato."],
        ["Quote (Orçamento)", "0-14",
         "Vinculado ao Deal. Controlado por total_de_orcamentos_* e hs_quote_esign_status. "
         "Sinaliza quando o ganho manual deve ser dado."],
        ["Ticket (Engajamento)", "0-5",
         "Source para Deal de Retração. canal_responsavel é copiado do ticket para o Deal."],
    ],
)
set_col_widths(t41, [3.5, 3, 11])
doc.add_paragraph()

# 4.2 Propriedades Críticas
add_heading(doc, "4.2 Propriedades Críticas", level=2)

add_para(doc, "Propriedades de identidade e roteamento", bold=True, color_key="navy")
t42a = create_brand_table(doc,
    ["Propriedade (API)", "Label", "Tipo", "Papel"],
    [
        ["canal_responsavel", "Canal oficial", "enumeration",
         "Campo de roteamento principal. Em Expansão, define o aprovador. "
         "Em Aquisição, o campo de roteamento para aprovação é canal_implantacao."],
        ["canal_implantacao", "Canal implantação", "enumeration",
         "Usado para roteamento de aprovadores em Aquisição."],
        ["canal_legado_oficiais__legado", "Canal Legado", "enumeration",
         "Copiado da Empresa. Usado como chave de associação Deal→objeto Canal."],
        ["origem", "Origem", "enumeration",
         "Define a fila de round robin: Inside Sales, BackOffice ou Canal."],
        ["tipo_de_orcamento", "Tipo de orçamento", "enumeration",
         "Define o pipeline: Aquisição, Expansão ou Retração."],
        ["modalidade_de_aquisicao", "Modalidade de aquisição", "enumeration",
         "SaaS ou LU com DC. Determina cálculo financeiro e fluxo de aprovação."],
        ["detalhamento_orcamento", "Detalhamento orçamento", "enumeration",
         "Sub-tipo de retração/expansão (ex: redução temporária, conector, migração)."],
        ["id_sienge", "ID Sienge", "number",
         "ID do produto contratado. Não reutilizado entre produtos diferentes."],
        ["id_do_cliente", "ID do cliente", "number",
         "ID do cliente no sistema Sienge."],
    ],
)
set_col_widths(t42a, [5.5, 3.5, 2, 6.5])
doc.add_paragraph()

add_para(doc, "Propriedades de equipe (Round Robin)", bold=True, color_key="navy")
t42b = create_brand_table(doc,
    ["Propriedade (API)", "Label", "Papel"],
    [
        ["equipe_vendas__inside_sales__vendedor", "[Equipe] Vendas - Inside Sales - Vendedor",
         "Vendedor atribuído via round robin para Aquisição com origem Inside Sales."],
        ["equipe_vendas__backoffice__vendedor", "[Equipe] Vendas - BackOffice - Vendedor",
         "Vendedor atribuído via round robin para Aquisição com origem BackOffice."],
        ["equipe_canal__vendedor", "[Equipe] Canal - Vendedor",
         "Vendedor atribuído via round robin para Aquisição com origem Canal (parceiro)."],
        ["equipe_expansao__inside_sales__vendedor", "[Equipe] Expansão - Inside Sales - Vendedor",
         "Vendedor atribuído para Expansão via Inside Sales."],
        ["equipe_contratos__analista", "[Equipe] Contratos - Analista",
         "Analista de contratos atribuído via round robin após o envio para aprovação."],
        ["gatilho_de_atribuicao", "Gatilho de atribuição",
         "Dispara o workflow de round robin para reatribuição."],
    ],
)
set_col_widths(t42b, [6, 5, 6.5])
doc.add_paragraph()

add_para(doc, "Propriedades de contrato e assinatura", bold=True, color_key="navy")
t42c = create_brand_table(doc,
    ["Propriedade (API)", "Label", "Tipo", "Papel"],
    [
        ["negocio_enviado_para_ganho_manual", "Negócio enviado para ganho manual?", "enumeration",
         "Sinaliza que o financeiro foi notificado para dar o ganho. Setado pelo workflow de ganho."],
        ["data_assinatura_contrato_pai", "Data Assinatura Contrato Pai", "date",
         "Preenchida automaticamente no momento da assinatura de todos os signatários externos."],
        ["total_de_orcamentos_assinados", "Total de orçamentos assinados", "number",
         "Contador que dispara o workflow de ganho manual quando incrementado."],
        ["gatilho_para_ganho_automatico", "Gatilho para ganho automático", "enumeration",
         "Campo no Quote (0-14) que habilita o fluxo de ganho manual ao ser setado como Sim."],
        ["autorizar_passagem_para_contrato_assinado", "Autorizar passagem para Contrato Assinado?", "bool",
         "Trava de etapa no pipeline de Retração. Impede CS de mover sem autorização."],
        ["autorizar_passagem_para_vendido", "Autorizar passagem para Vendido?", "bool",
         "Trava de etapa para Vendido. Preenchida pelo aprovador."],
        ["gatilho_de_contrato_assinado", "Gatilho de contrato assinado", "enumeration",
         "Dispara processos pós-assinatura (Freshdesk, CS)."],
        ["nr_contrato", "Número do Contrato", "string",
         "Gerado pelo processo 2.1 após o ganho. Referência para o processo 3.0 e 5.0."],
    ],
)
set_col_widths(t42c, [5.5, 4.5, 2, 5.5])
doc.add_paragraph()

add_para(doc, "Propriedades de SLA e próximos passos", bold=True, color_key="navy")
t42d = create_brand_table(doc,
    ["Propriedade (API)", "Label", "Tipo", "Papel"],
    [
        ["proximos_passos", "Próximos passos", "string",
         "Preenchido pelo vendedor. Gera tarefa automática com lembrete."],
        ["data_da_proxima_etapa", "Data dos próximos passos", "date",
         "Data do lembrete. Dispara criação de tarefa quando preenchida junto com proximos_passos."],
        ["previsao_de_fechamento", "Previsão de fechamento", "date",
         "Exibida no card de negócios. Distinta da data de fechamento real (closedate)."],
        ["closedate", "Close Date", "datetime",
         "Preenchida apenas no fechamento efetivo (ganho ou perda). Não deve ser confundida com previsão."],
    ],
)
set_col_widths(t42d, [5.5, 4, 2, 6])
doc.add_paragraph()

add_para(doc, "Pipelines de Deal", bold=True, color_key="navy")
add_para(doc, "Aquisição (ID: 795178644) e Expansão (ID: 839459507):", bold=False, color_key="dark_gray")
t42e_aq = create_brand_table(doc,
    ["Etapa", "Observação"],
    [
        ["Oportunidade Identificada", "Etapa inicial."],
        ["Oportunidade Qualificada", ""],
        ["Agendamento", "SLA: alerta ao vendedor após X dias sem avanço."],
        ["Solução", "Preencher Próximos Passos obrigatório ao avançar de Agendamento para aqui."],
        ["Negociação", "Preencher tipo_de_orcamento e modalidade_de_aquisicao."],
        ["Contrato", "Envio para aprovação e assinatura. Workflow de ganho manual ativo."],
        ["Upsell", "Etapa de ganho para Expansão."],
        ["Vendido", "Etapa de ganho para Aquisição. Dispara processo 2.1."],
        ["Perdido", "Etapa de perda."],
        ["Descartado", "Etapa de descarte."],
    ],
)
set_col_widths(t42e_aq, [5, 12.5])
doc.add_paragraph()

add_para(doc, "Retração (ID: 842042121):", bold=False, color_key="dark_gray")
t42e_ret = create_brand_table(doc,
    ["Etapa", "Observação"],
    [
        ["Solicitação", "Criado automaticamente pelo playbook de churn."],
        ["Diagnóstico e Solução", "SLA: 12 dias (corridos) entre aviso prévio e envio para aprovação."],
        ["Em Formalização", "Orçamento sendo gerado automaticamente. Deal fica aqui enquanto o código executa."],
        ["Orçamento Enviado", "Orçamento gerado. SLA: 6 dias para aprovação do time de contratos."],
        ["Contrato Assinado", "Requer autorizar_passagem_para_contrato_assinado = Sim. SLA: 6 dias para assinatura."],
        ["Retração Realizada", "Etapa de ganho. SLA: 1 dia para registrar."],
        ["Descarte", "Etapa de descarte. Orçamento é expirado automaticamente."],
    ],
)
set_col_widths(t42e_ret, [5, 12.5])
doc.add_paragraph()

# 4.3 Workflows
add_heading(doc, "4.3 Workflows Envolvidos", level=2)

add_para(doc, "Cluster: Movimentação de Etapas", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="Nx. [Movimentação] Tipo Orçamento / Modalidade Aquisição → Move etapa",
    id_wf="1753092609",
    objetivo=(
        "Mover o Deal para a etapa de Contrato correta com base no tipo de orçamento e na modalidade "
        "de aquisição preenchidos."
    ),
    objeto="Deal (0-3)",
    disparo=(
        "Deal no pipeline de Aquisição (795178644) com tipo_de_orcamento ou modalidade_de_aquisicao "
        "preenchidos. Re-inscrição ativa."
    ),
    acoes=(
        "Branch por tipo_de_orcamento:\n"
        "  Aquisição → sub-branch por modalidade_de_aquisicao:\n"
        "    SaaS → move para etapa de Contrato SaaS.\n"
        "    LU com DC → move para etapa de Contrato LU.\n"
        "  Expansão → lógica similar com etapas de Expansão.\n"
        "  Retração → move para etapa Em Formalização.\n"
        "Workflows análogos existem para Expansão e Retração."
    ),
    dependencias="tipo_de_orcamento e modalidade_de_aquisicao devem estar preenchidos no Deal."
)

add_para(doc, "Cluster: Round Robin e Atribuição", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="[Round Robin] Deal Aquisição com Origem conhecida → Define Equipes responsáveis",
    id_wf="1798412036",
    objetivo=(
        "Distribuir Deals do pipeline de Aquisição entre os membros da equipe de vendas "
        "de forma rotativa, conforme a origem do Deal."
    ),
    objeto="Deal (0-3)",
    disparo=(
        "Deal no pipeline de Aquisição com origem e canal_responsavel conhecidos. Re-inscrição ativa "
        "via propriedade gatilho_de_atribuicao."
    ),
    acoes=(
        "Branch por origem:\n"
        "  Origem = 'Inside Sales' → round robin para equipe_vendas__inside_sales__vendedor.\n"
        "  Origem = 'BackOffice' → round robin para equipe_vendas__backoffice__vendedor.\n"
        "  Origem = 'Canal' → round robin para equipe_canal__vendedor.\n"
        "Workflow com ~104 ações (payload grande) — distribui entre múltiplos membros de cada equipe."
    ),
    dependencias="origem e canal_responsavel devem estar preenchidos."
)

add_workflow_block(
    doc,
    nome="[Round Robin] Deal Aquisição com Canal → Define Equipes responsáveis",
    id_wf="1784845853",
    objetivo=(
        "Complementar o round robin de Aquisição para Deals com canal_responsavel específico "
        "e sem origem definida."
    ),
    objeto="Deal (0-3)",
    disparo="Deal no pipeline de Aquisição com canal_responsavel conhecido e sem equipe atribuída.",
    acoes="Round robin por canal_responsavel para a equipe de vendas correspondente.",
    dependencias="canal_responsavel deve estar preenchido."
)

add_para(doc, "Cluster: Canal Oficial", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="[Associação] Negócio Criado & Canal Oficial conhecido → Associa ao Canal respectivo",
    id_wf="1793578765",
    objetivo=(
        "Copiar o canal_legado do Canal associado para o Deal e criar a associação "
        "Deal → objeto Canal (2-54226027)."
    ),
    objeto="Deal (0-3)",
    disparo=(
        "canal_responsavel preenchido OU canal_legado_oficiais__legado preenchido no Deal. "
        "Re-inscrição quando canal_responsavel é atualizado."
    ),
    acoes=(
        "1. Branch: canal_responsavel preenchido → copia canal_legado do objeto Canal associado "
        "para canal_legado_oficiais__legado do Deal.\n"
        "2. Aguarda 1 minuto.\n"
        "3. Cria associação Deal → Canal (2-54226027) por matching: "
        "Deal.canal_legado_oficiais__legado = Canal.canais_oficiais_e_legados.\n"
        "4. Copia canal_responsavel do Canal para o Deal.\n"
        "5. Seta gatilho_canal_associado_ao_deal = Sim."
    ),
    dependencias="Objeto Canal (2-54226027) deve ter a propriedade canais_oficiais_e_legados preenchida."
)

add_para(doc, "Cluster: Ganho Manual", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="Nx. [Fluxo de envio para ganho manual] — Aquisição",
    id_wf="1791344691",
    objetivo=(
        "Notificar o financeiro para dar o ganho manual quando todos os orçamentos do Deal "
        "estão assinados e o gatilho de ganho automático está ativo."
    ),
    objeto="Deal (0-3)",
    disparo=(
        "Deal na etapa 1166928285 (Contrato/enviado para assinatura) com Quote associado: "
        "hs_quote_esign_status = SIGNED, hs_quote_status = PUBLISHED, "
        "gatilho_para_ganho_automatico = Sim. "
        "Re-inscrição quando total_de_orcamentos_assinados é atualizado."
    ),
    acoes=(
        "1. Seta data_assinatura_contrato_pai = timestamp de execução.\n"
        "2. Seta negocio_enviado_para_ganho_manual = Sim.\n"
        "3. Envia alerta ao owner: 'Minutas assinadas! Dar ganho no negócio.'\n"
        "4. Envia alerta ao time: 'Negócio de Aquisição foi enviado para o ganho manual!'\n"
        "5. Branch de verificação: se ainda há Quotes com PENDING_SIGNATURE, envia alerta "
        "'Negócio ganho mas com orçamentos pendentes'.\n"
        "Workflow análogo existe para Expansão."
    ),
    dependencias=(
        "Requer Deal na etapa de Contrato. "
        "Deal deve ter total_de_orcamentos_assinados preenchido (contador). "
        "Quote deve ter gatilho_para_ganho_automatico = Sim."
    )
)

add_para(doc, "Cluster: SLA e Notificações", bold=True, size=12, color_key="primary_black")
add_para(doc,
    "Workflows de SLA existem para cada etapa dos pipelines de Aquisição, Expansão e Retração. "
    "Cada workflow monitora o tempo de permanência em uma etapa e dispara notificação ao responsável "
    "quando o prazo é excedido. Os SLAs abaixo são do pipeline de Retração (ticket 86b8m3dkx). "
    "Os SLAs de Aquisição e Expansão não foram formalizados.",
    size=10.5, indent=True
)
t_sla = create_brand_table(doc,
    ["Etapa", "Prazo", "Início do período", "Fim do período", "Alerta"],
    [
        ["Solicitação / Diagnóstico / Solução", "12 dias corridos",
         "Data de início do aviso prévio", "Envio do contrato para aprovação",
         "Customer Success"],
        ["Formalização — aprovação do contrato", "6 dias",
         "Recebimento da aprovação do contrato", "Aprovação do financeiro",
         "Customer Success, Financeiro"],
        ["Formalização — assinatura do cliente", "6 dias",
         "Aprovação do financeiro", "Assinatura do cliente",
         "Customer Success"],
        ["Formalização — processamento financeiro", "5 dias",
         "Assinatura do cliente", "Processamento financeiro",
         "Financeiro"],
        ["Final — registrar ganho", "1 dia",
         "Assinatura completa", "Registro em \"Retração Realizada\"",
         "—"],
    ],
)
set_col_widths(t_sla, [4, 2, 4, 4, 3.5])
doc.add_paragraph()

add_para(doc, "Cluster: Cópia de Canal para Deal de Retração", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="[Cópia] Canal do Ticket de Retração → Deal de Retração",
    id_wf="(ver ticket 86b9yr5vj / 86ba8b0vy)",
    objetivo=(
        "Copiar canal_responsavel do ticket de engajamento de origem para o Deal de Retração "
        "automaticamente na criação."
    ),
    objeto="Deal (0-3)",
    disparo="Deal de Retração criado sem canal_responsavel preenchido e com ticket de engajamento associado.",
    acoes="Copia canal_responsavel do ticket associado para o Deal de Retração.",
    dependencias="Ticket de engajamento deve estar associado ao Deal de Retração."
)

add_para(doc, "Cluster: Conversão LT Diário", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="[Conversão] LT Diário texto → número",
    id_wf="(ver ticket 86ba8k6gu)",
    objetivo="Converter o valor do campo LT Diário (texto) para o campo LT Diário (número) no Deal.",
    objeto="Deal (0-3)",
    disparo="Deal com campo LT Diário (texto) preenchido.",
    acoes="Copia e converte o valor do campo de texto para o campo numérico correspondente.",
    dependencias="Campos LT Diário (texto) e LT Diário (número) devem existir no objeto Deal."
)

add_para(doc, "Cluster: Próximos Passos", bold=True, size=12, color_key="primary_black")
add_workflow_block(
    doc,
    nome="[Tarefa] Próximos Passos preenchidos → cria tarefa com lembrete",
    id_wf="(ver ticket 86ba4qgq7)",
    objetivo="Criar tarefa automática com lembrete quando o vendedor preenche Próximos Passos e a data.",
    objeto="Deal (0-3)",
    disparo="proximos_passos e data_da_proxima_etapa preenchidos no Deal.",
    acoes=(
        "1. Cria tarefa no HubSpot associada ao Deal com o conteúdo de proximos_passos.\n"
        "2. Define o lembrete para a data_da_proxima_etapa.\n"
        "Nota: o workflow não apaga o campo proximos_passos após criar a tarefa."
    ),
    dependencias="proximos_passos e data_da_proxima_etapa devem estar preenchidos simultaneamente."
)

# 4.4 Cards e Customizações
add_heading(doc, "4.4 Cards e Customizações", level=2)
add_para(doc,
    "O processo 2.0 utiliza os seguintes recursos customizados no HubSpot:"
)
doc.add_paragraph()

add_para(doc, "Playbooks", bold=True, color_key="navy")
for item in [
    "Passagem de Vendas (GO): playbook preenchido pelo vendedor ao avançar o Deal. Registra informações de handover para CS.",
    "Formulário de Entrada: playbook utilizado para qualificação inicial da oportunidade.",
    "Cancelamento ou Redução: playbook preenchido pelo CS para iniciar o processo de Retração. Dispara criação automática do Deal de Retração e do ticket de retração.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc, "Filtros e visões personalizadas", bold=True, color_key="navy")
for item in [
    "Filtro 'Contratos Aguardando Aprovação do Time de Contratos': exibe Deals pendentes de aprovação.",
    "Filtro 'Contratos Reprovados': exibe Deals cujo contrato foi recusado (automática ou manualmente).",
    "Filtro 'Contratos aprovados aguardando assinatura': exibe Deals aprovados mas não assinados.",
    "Filtro 'Contratos assinados aguardando o ganho': exibe Deals com contrato assinado aguardando ação do financeiro.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc, "Propriedade no card de negócios", bold=True, color_key="navy")
add_para(doc,
    "A previsão de fechamento (previsao_de_fechamento) foi incluída no resumo do card de negócios "
    "nos pipelines de Aquisição e Expansão (ticket 86ba3tjv7). "
    "A data de fechamento real (closedate) não deve aparecer antes do ganho/perda efetivo.",
    size=10.5, indent=True
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 5. RISCOS E DEPENDÊNCIAS
# ════════════════════════════════════════════════════════
add_heading(doc, "5. Riscos e Dependências", level=1)

t5 = create_brand_table(doc,
    ["Risco / Dependência", "Impacto", "Mitigação atual"],
    [
        ["canal_responsavel alterado indevidamente por fluxo automático",
         "Contrato vai para aprovador errado. Ocorreu em Expansão (ticket 86ba43rfc): fluxo removeu o "
         "canal_responsavel preenchido manualmente pelo vendedor.",
         "Fluxo corrigido na estabilização. Monitorar re-ocorrências."],
        ["canal_responsavel não copiado do ticket de retração para o Deal de Retração",
         "Deal de Retração sem canal → cai para aprovador errado. Ocorreu múltiplas vezes "
         "(tickets 86b9yr5vj, 86ba8b0vy).",
         "Fluxo de cópia corrigido. Monitorar novos Deals de Retração."],
        ["Signatários obrigatórios ausentes no contrato",
         "Contrato reprovado automaticamente, gerando retrabalho para o vendedor.",
         "Workflow de verificação reprova e notifica. Depende de treinamento do time comercial."],
        ["Deal de Aquisição criado sem gerar negócio (prospecção do canal)",
         "Fluxo de criação automática de Deal a partir de Lead não disparou (ticket 86ba50uw0). "
         "Canal criou o contato e avançou etapas sem Deal associado.",
         "Fluxo corrigido. Canal orientado a criar Deal manualmente em caso de falha."],
        ["Valor do negócio incorreto ou com atraso de atualização",
         "Valores errados confundem vendedores e time de contratos. Sobrescrita por fluxo "
         "após ajuste manual (ticket 86b9xz7vk).",
         "Ticket marcado como done. Identificado como problema parcial do HubSpot. "
         "Monitorar ocorrências."],
        ["Playbook de Inadimplência não cria ticket de retração",
         "CS não consegue iniciar o processo de retração por inadimplência via playbook "
         "(ticket 86bacyj4b — blocked).",
         "Sem mitigação automática. Ponto a Validar PV-01."],
        ["Expiração de orçamento de retração pós descarte falhando",
         "Orçamentos ficam ativos mesmo após descarte do Deal de Retração "
         "(ticket 86bacv6ap — blocked).",
         "Sem mitigação confirmada. Ponto a Validar PV-02."],
        ["ID Sienge incorreto vs. Vtiger",
         "Clientes com ID Sienge diferente do ID no Vtiger causam problemas de integração e "
         "operação (tickets 86b9xkx7h, 86ba66trp, 86ba62uc0).",
         "Correção em massa realizada. Regra documentada: sem reaproveitamento entre produtos."],
    ],
)
set_col_widths(t5, [5, 5.5, 7])
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 6. MATERIAIS DE APOIO
# ════════════════════════════════════════════════════════
add_heading(doc, "6. Materiais de Apoio", level=1)

t6mat = create_brand_table(doc,
    ["Tipo", "Título / ID", "Localização / Status"],
    [
        ["Ticket ClickUp", "86bad8x7k — Ajuste no fluxo de aprovação (separação aprovadores)", "done"],
        ["Ticket ClickUp", "86bacyj4b — Playbook de Inadimplência não abre ticket de retração", "blocked"],
        ["Ticket ClickUp", "86bacv6ap — Expiração de orçamento de retração pós descarte", "blocked"],
        ["Ticket ClickUp", "86ba8k6gu — Fluxo de correção de campo LT Diário", "done"],
        ["Ticket ClickUp", "86ba8b0vy — Troca de canal não gerou ticket e alerta contratos", "done"],
        ["Ticket ClickUp", "86ba4qgq7 — Ajuste na criação de tarefas de Próximos Passos", "done"],
        ["Ticket ClickUp", "86ba43rfc — Alteração indevida de canal oficial por fluxo automático", "done"],
        ["Ticket ClickUp", "86ba3tjv7 — Inserir previsão de fechamento no card de negócios", "done"],
        ["Ticket ClickUp", "86ba01bah — ID Sienge — regra de não reaproveitamento entre produtos", "done"],
        ["Ticket ClickUp", "86b9wkb3f — Obrigar signatários Maria Madalena e Giovani Amaral", "done"],
        ["Ticket ClickUp", "86b9wkkkz — Alerta cliente assinou com signatários internos pendentes", "done"],
        ["Ticket ClickUp", "86b9xz7vk — Valor do negócio com problemas de atraso e inserção incorreta", "done"],
        ["Ticket ClickUp", "86b8m3dkx — Criação dos SLAs para Retração ou Churn", "done"],
        ["Ticket ClickUp", "86b8m3dkx — Cópia de propriedades do playbook para ticket de Retração", "done"],
        ["Workflow HubSpot", "Movimentação Tipo/Modalidade Aquisição → Move etapa", "ID 1753092609"],
        ["Workflow HubSpot", "Round Robin Deal Aquisição com Origem conhecida", "ID 1798412036"],
        ["Workflow HubSpot", "Round Robin Deal Aquisição com Canal", "ID 1784845853"],
        ["Workflow HubSpot", "Associação Negócio & Canal Oficial → Associa ao Canal", "ID 1793578765"],
        ["Workflow HubSpot", "Fluxo de envio para ganho manual — Aquisição", "ID 1791344691"],
    ],
)
set_col_widths(t6mat, [3.5, 9, 5])
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 7. PONTOS A VALIDAR
# ════════════════════════════════════════════════════════
add_heading(doc, "7. Pontos a Validar", level=1)

add_callout(doc, "warning",
    "Ação requerida — Vinícius Vieira Braz",
    "Os pontos abaixo não foram confirmados durante a documentação e requerem validação do responsável "
    "(Vinícius Vieira Braz) antes da publicação definitiva deste documento."
)
doc.add_paragraph()

t7 = create_brand_table(doc,
    ["ID", "Ponto", "Detalhe"],
    [
        ["PV-01", "Playbook de Inadimplência não cria ticket de retração (blocked)",
         "O ticket 86bacyj4b está bloqueado: o playbook de cancelamento e redução com tipo = "
         "Inadimplência não está criando o ticket de retração. "
         "Confirmar: (a) o fluxo foi investigado e corrigido após a documentação? "
         "(b) qual é o estado atual do workflow responsável por essa criação? "
         "(c) se não corrigido, como o CS deve proceder enquanto o bug persiste?"],
        ["PV-02", "Expiração de orçamento de retração pós descarte (blocked)",
         "O ticket 86bacv6ap está bloqueado: o processo de expirar orçamento após o descarte do "
         "Deal de Retração não está funcionando para alguns casos. "
         "Confirmar: (a) o workflow de expiração foi corrigido? "
         "(b) quais casos específicos falham? "
         "(c) existe fallback manual para expirar orçamentos órfãos?"],
        ["PV-03", "Lógica completa do Round Robin — equipes e critério de distribuição",
         "Os workflows de Round Robin (1784845853 e 1798412036) têm payloads grandes e não foram "
         "lidos na íntegra. A estrutura geral está documentada (Inside Sales / BackOffice / Canal). "
         "Confirmar: (a) quais membros compõem cada fila de round robin atualmente? "
         "(b) existe round robin para Expansão? (c) existe mecanismo para re-atribuição quando "
         "um vendedor está ausente?"],
        ["PV-04", "Tickets de onboarding não enviando formulário de saída (blocked)",
         "O ticket 86ba5rm2c informa que tickets de onboarding do Sienge Plataforma migrados "
         "não disparam o formulário de saída ao mover para 'Cliente Ativado'. "
         "Confirmar: (a) esse problema foi corrigido no processo de CS (6.0)? "
         "(b) afeta apenas tickets migrados ou também os novos?"],
        ["PV-05", "Ativação dos fluxos de atribuição automática de CS (to do)",
         "O ticket 86b93aev6 ([Processos] Ativação dos fluxos de atribuição automática de CS) "
         "está com status 'to do'. "
         "Confirmar: (a) quando será ativado? "
         "(b) qual é o impacto atual de não ter essa atribuição automática?"],
        ["PV-06", "Campos Sistema e Módulo não populados no Ticket de Engajamento (to do)",
         "O ticket 86b9yr5vj aponta que campos (Produtos Starian) e Módulo estão vazios em "
         "praticamente todos os tickets de engajamento. "
         "Confirmar: (a) esses campos são necessários para o processo 2.0? "
         "(b) o preenchimento é manual ou deveria ser automático?"],
        ["PV-07", "SLAs do pipeline de Aquisição — prazos exatos por etapa",
         "Os SLAs de Retração foram documentados detalhadamente (ticket 86b8m3dkx). "
         "Os SLAs de Aquisição e Expansão não foram documentados com os prazos exatos. "
         "Confirmar: (a) quais são os prazos de SLA por etapa nos pipelines de Aquisição e Expansão? "
         "(b) os alertas de SLA estão ativos para todas as etapas?"],
    ],
)
set_col_widths(t7, [1.5, 4.5, 11.5])
doc.add_paragraph()

doc.save(OUTPUT_PATH)
print(f"Documento gerado: {OUTPUT_PATH}")
