"""
Gerador de documentação para o processo 8.1 - Integrações Oracle
Executar: python3 gerar_documento.py
"""

import sys
from pathlib import Path

# ── nexforce-brand helpers ──────────────────────────────────────────────────
SKILL_DIR = Path("/home/hugo-zanni/Nexforce/Projects/handover-sienge/.claude/skills/nexforce-brand")
sys.path.insert(0, str(SKILL_DIR / "references"))

import python_docx_helpers as _brand
_brand.nexforce_assets = lambda: SKILL_DIR / "assets"

from python_docx_helpers import (
    apply_nexforce_theme,
    insert_brand_header,
    insert_brand_footer,
    create_brand_table,
    add_callout,
    NEXFORCE_COLORS,
    logo_path,
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "8.1 - Integrações Oracle.docx"


def _rgb(key: str) -> RGBColor:
    return RGBColor.from_string(NEXFORCE_COLORS[key])


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Lato"
    if level == 1:
        run.font.color.rgb = _rgb("primary_black")
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = _rgb("navy")
        run.font.size = Pt(13)
        run.font.bold = True
    else:
        run.font.color.rgb = _rgb("dark_gray")
        run.font.size = Pt(11)
        run.font.bold = True
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color_key="dark_gray", indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Lato"
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color_key)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    for run in p.runs:
        run.font.name = "Lato"
        run.font.size = Pt(10.5)
        run.font.color.rgb = _rgb("dark_gray")
    return p


def add_workflow_block(doc, nome, id_wf, objeto, disparo, acoes, dependencias=None, status="Ativo"):
    p = doc.add_paragraph()
    run = p.add_run(f"Workflow: {nome}")
    run.bold = True
    run.font.name = "Lato"
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb("navy")

    rows_data = [
        ("ID", id_wf),
        ("Status", status),
        ("Objeto HubSpot", objeto),
        ("Disparo", disparo),
        ("Ações principais", acoes),
    ]
    if dependencias:
        rows_data.append(("Dependências", dependencias))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        cell_l = table.rows[i].cells[0]
        cell_l.text = label
        _set_cell_bg(cell_l, NEXFORCE_COLORS["callout_info_bg"])
        for para in cell_l.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Lato"
                run.font.size = Pt(9.5)
                run.font.color.rgb = _rgb("navy")

        cell_v = table.rows[i].cells[1]
        cell_v.text = value
        bg = NEXFORCE_COLORS["white"] if i % 2 == 0 else NEXFORCE_COLORS["near_white"]
        _set_cell_bg(cell_v, bg)
        for para in cell_v.paragraphs:
            for run in para.runs:
                run.font.name = "Lato"
                run.font.size = Pt(9.5)
                run.font.color.rgb = _rgb("dark_gray")

    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(13)
    doc.add_paragraph()


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
apply_nexforce_theme(doc)

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

insert_brand_header(doc, title="8.1 — Integrações Oracle", period="Junho 2026")
insert_brand_footer(doc, confidentiality="Uso interno  ·  Nexforce Services  ·  v1.0")

# ════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Cm(8)
r = p_sub.add_run("Documentação de Processos — HubSpot RaaS")
r.font.name = "Lato"
r.font.size = Pt(12)
r.font.color.rgb = _rgb("mid_gray")

doc.add_paragraph()

p_title = doc.add_heading("8.1 — Integrações Oracle", level=0)
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in p_title.runs:
    run.font.name = "Lato"
    run.font.color.rgb = _rgb("primary_black")
    run.font.size = Pt(26)

doc.add_paragraph()

meta_table = create_brand_table(
    doc,
    headers=["Campo", "Valor"],
    rows=[
        ["Cliente", "Sienge"],
        ["Responsável pelo processo", "Vinicius Vanoni"],
        ["Empresa responsável", "Nexforce Services"],
        ["Data de documentação", "15 de junho de 2026"],
        ["Versão", "1.0"],
        ["Status do processo", "Estabilização pós Go Live"],
    ],
)
set_col_widths(meta_table, [6, 10])

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SUMÁRIO (manual)
# ════════════════════════════════════════════════════════
add_heading(doc, "Sumário", level=1)
for item in [
    "1. Objetivo",
    "2. Contexto de Negócio",
    "3. Visão Funcional",
    "   3.1 Fluxo Operacional",
    "   3.2 Regras de Negócio",
    "   3.3 Critério de Validação",
    "4. Visão Técnica",
    "   4.1 Objetos HubSpot Envolvidos",
    "   4.2 Propriedades Críticas",
    "   4.3 Workflows Envolvidos",
    "   4.4 Integração Oracle",
    "   4.5 Cards e Customizações",
    "5. Riscos e Dependências",
    "6. Materiais de Apoio",
    "7. Pontos a Validar",
]:
    add_para(doc, item, size=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. OBJETIVO
# ════════════════════════════════════════════════════════
add_heading(doc, "1. Objetivo", level=1)
add_para(doc,
    "Manter o Grupo de Contrato no HubSpot sincronizado com os valores financeiros reais do Oracle ERP do Sienge, "
    "garantindo que os itens de linha do Grupo de Contrato reflitam o MRR por tipo de serviço contratado e, quando "
    "aplicável, por módulo/sistema."
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 2. CONTEXTO DE NEGÓCIO
# ════════════════════════════════════════════════════════
add_heading(doc, "2. Contexto de Negócio", level=1)

add_para(doc,
    "O Oracle é o sistema de faturamento e gestão de contratos do Sienge. Os valores reais de cobrança de cada cliente "
    "— SaaS, Manutenção, APIs, NFe, Data Center, conectores, entre outros — são calculados e armazenados no Oracle e "
    "precisam estar disponíveis no HubSpot para que as equipes de CS, BackOffice e gestão possam:"
)
for item in [
    "Acompanhar o MRR real de cada cliente no Grupo de Contrato;",
    "Comparar o valor contratado (comercial) com o valor faturado (Oracle);",
    "Detectar divergências de valor que sinalizam renegociações, aditivos não processados ou erros de faturamento;",
    "Alimentar KPIs de receita e churn com dados financeiros acurados.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(doc,
    "O Grupo de Contrato é o objeto HubSpot que centraliza todos os contratos de um mesmo CNPJ sob um único registro. "
    "A integração Oracle atualiza diariamente as propriedades de valor (prog_qtd_*) desse objeto e dispara a reconstrução "
    "automática de seus itens de linha."
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 3. VISÃO FUNCIONAL
# ════════════════════════════════════════════════════════
add_heading(doc, "3. Visão Funcional", level=1)

# 3.1 Fluxo Operacional
add_heading(doc, "3.1 Fluxo Operacional", level=2)
add_para(doc, "O processo segue a cadência diária abaixo:")
doc.add_paragraph()

steps = [
    ("1", "Oracle disponibiliza dados diários",
     "O Oracle ERP do Sienge expõe um endpoint REST com os valores de contrato de cada cliente para a data atual. "
     "Os dados incluem número do contrato, tipo de serviço (lin_servico / lin_servico_desc) e valor faturado (prog_qtd)."),
    ("2", "Middleware autentica e busca os dados",
     "O workflow [DEV] Integração Oracle (Grupo de Contrato) autentica no Oracle via usuário e senha (secrets ORACLE_USERNAME / "
     "ORACLE_PASSWORD), obtém um token e consulta o endpoint de contratos passando a data do dia no formato DDMMYYYY."),
    ("3", "Dados são gravados nas propriedades prog_qtd_*",
     "Os valores retornados pelo Oracle são mapeados para as propriedades prog_qtd_* do Grupo de Contrato correspondente. "
     "O ID do Grupo de Contrato é derivado do nr_contrato Oracle usando o prefixo antes do primeiro ponto, underline ou hífen."),
    ("4", "Workflow de Split é disparado automaticamente",
     "Qualquer alteração em uma propriedade prog_qtd_* ou na propriedade sistema dispara o re-enrollment no workflow "
     "[01. Integração Oracle - Grupo de Contrato] Split de serviços em itens de linha (v03)."),
    ("5", "Itens de linha existentes são arquivados",
     "O workflow executa um batch archive (delete-all) de todos os itens de linha atualmente associados ao Grupo de Contrato. "
     "Não há atualização item a item — todos são apagados e recriados a cada execução."),
    ("6", "Novos itens de linha são criados com os valores atualizados",
     "O workflow determina se o caso requer split proporcional por sistema (Com Split) ou criação direta por serviço (Sem Split) "
     "e cria os itens de linha no HubSpot associados ao Grupo de Contrato."),
    ("7", "Campos financeiros consolidados são recalculados",
     "Após a atualização dos prog_qtd_*, workflows secundários recalculam os campos Valor Serviço, Valor SaaS, "
     "Valor Manutenção Mensal (LU) e Valor Data Center (LU) no Grupo de Contrato."),
]

for num, titulo, descricao in steps:
    p = doc.add_paragraph()
    r_num = p.add_run(f"Passo {num}: ")
    r_num.bold = True
    r_num.font.name = "Lato"
    r_num.font.color.rgb = _rgb("navy")
    r_num.font.size = Pt(11)
    r_titulo = p.add_run(titulo)
    r_titulo.bold = True
    r_titulo.font.name = "Lato"
    r_titulo.font.size = Pt(11)
    r_titulo.font.color.rgb = _rgb("primary_black")
    add_para(doc, descricao, size=10.5, indent=True)
    doc.add_paragraph()

# 3.2 Regras de Negócio
add_heading(doc, "3.2 Regras de Negócio", level=2)

regras = [
    ("Regra 1 — Identificação do Grupo de Contrato",
     "O ID do Grupo de Contrato no HubSpot (nr_contrato_hs) é derivado do número de contrato Oracle aplicando "
     "a regex ^[^._]+ — captura tudo antes do primeiro ponto (.), underline (_) ou hífen (-). "
     'Exemplo: nr_contrato "12345.001" gera grupo_id "12345".'),
    ("Regra 2 — Estratégia delete-all + recreate",
     "A cada execução do workflow de split, todos os itens de linha existentes do Grupo de Contrato são arquivados "
     "(batch archive) antes de criar os novos. Não há tentativa de atualizar item a item. "
     "Isso evita inconsistências mas significa que, durante a execução (~segundos), o GC fica temporariamente sem itens."),
    ("Regra 3 — Split de SaaS / Manutenção por módulo",
     "Quando a propriedade sistema está preenchida e prog_qtd_saas ou prog_qtd_manutencao tem valor, "
     "o valor total é distribuído proporcionalmente entre os sistemas/módulos contratados. "
     "O peso de cada módulo é determinado pelo preço cadastrado no catálogo HubSpot (campo hs_price_brl do produto). "
     "Se dois módulos têm preços R$100 e R$400, o split é 20%/80% do valor total."),
    ("Regra 4 — Contrato misto SaaS + LU (Manutenção)",
     "Quando o cliente tem tanto prog_qtd_saas quanto prog_qtd_manutencao preenchidos: "
     "(a) o split proporcional por módulos usa o valor de prog_qtd_saas; "
     "(b) é criado um item adicional 'Manutenção Mensal' com o valor de prog_qtd_manutencao. "
     "Os itens de módulo são do tipo SaaS; o item de manutenção é separado."),
    ("Regra 5 — Sem sistemas (campo sistema vazio)",
     "Quando a propriedade sistema está vazia, é criado um único item genérico: "
     "'SaaS' (hs_sku: 44918051665) com o valor de prog_qtd_saas, "
     "ou 'Manutenção' (hs_sku: 39980447525) com o valor de prog_qtd_manutencao, conforme disponível."),
    ("Regra 7 — Split de Conectores",
     "O valor total de prog_qtd_connectors é distribuído proporcionalmente entre os produtos cujo nome começa com "
     "'Conector' ou 'Connectors', usando o preço do catálogo HubSpot como peso — mesma lógica do split de módulos."),
    ("Regra 8 — Anti-concorrência (execuções simultâneas)",
     "O workflow usa dois mecanismos para evitar que múltiplas alterações simultâneas em prog_qtd_* causem execuções "
     "paralelas e itens de linha duplicados: "
     "(a) random delay de 1 a 19 segundos no início de cada execução; "
     "(b) A/B test com 3 branches — uma executa imediatamente, outra aguarda 1 min, outra aguarda 2 min. "
     "Isso distribui o tempo de início das execuções concorrentes."),
    ("Regra 9 — Supressão de GCs cancelados",
     "Grupos de Contrato com a propriedade status = 'Cancelado' são excluídos do enrollment. "
     "O workflow não processa GCs de contratos encerrados."),
    ("Regra 10 — Itens de linha do Negócio (Deal) não são afetados",
     "O fluxo de integração Oracle opera exclusivamente sobre os itens de linha do Grupo de Contrato. "
     "Os itens de linha do Negócio (Deal) nunca devem ser alterados por este processo."),
    ("Regra 11 — Sem Split (quando saas/manutencao/connectors estão ausentes)",
     "Quando nenhum dos valores prog_qtd_saas, prog_qtd_manutencao ou prog_qtd_connectors está preenchido, "
     "o workflow executa o caminho 'Sem Split': cria itens de linha para todos os prog_qtd_* que tiverem valor, "
     "usando lin_servico_desc_* como nome do item e prog_qtd_* como preço. "
     "Neste caminho, os SKUs dos produtos conhecidos são aplicados via tabela estática."),
]

for titulo, texto in regras:
    add_para(doc, titulo, bold=True, color_key="primary_black")
    add_para(doc, texto, size=10.5, indent=True)
    doc.add_paragraph()

# Regra 6 — tabela de serviços fixos
add_para(doc, "Regra 6 — Serviços com valor fixo (sem split por módulo)", bold=True, color_key="primary_black")
add_para(doc,
    "Os seguintes serviços não passam por split proporcional — recebem o valor diretamente do prog_qtd correspondente:",
    size=10.5, indent=True
)
t6 = create_brand_table(doc,
    ["Serviço HubSpot", "Propriedade Oracle utilizada"],
    [
        ["APIs (Start/Essencial/Special/Ultimate/Enterprise)", "prog_qtd_apis"],
        ["NFe / NFE até 9 usuários / NFE acima de 10 usuários", "prog_qtd_nota_fiscal_eletronica"],
        ["NFs (Nota Fiscal de Serviço)", "prog_qtd_nota_fiscal_de_servico"],
        ["Recepção de CTe", "prog_qtd_emissao_de_conhecimento_de_transporte_eletronico_cte"],
        ["Base Teste (Sem/Com Anexo, Custo fixo)", "prog_qtd_base_de_testes"],
        ["E-Custos", "prog_qtd_ecustos_integracao"],
        ["Data Center LU", "prog_qtd_locacao_de_data_center"],
    ],
)
set_col_widths(t6, [10, 8])
doc.add_paragraph()

# 3.3 Critério de Validação
add_heading(doc, "3.3 Critério de Validação", level=2)
add_para(doc, "O processo foi executado corretamente quando:")
for item in [
    "Os itens de linha do Grupo de Contrato refletem os valores do Oracle para a data do dia.",
    "O total dos itens de linha (sum dos preços) corresponde ao MRR calculado pelo Oracle para o cliente.",
    "Para contratos SaaS com sistemas, cada módulo contratado tem um item de linha proporcional ao seu peso.",
    "Não há itens de linha duplicados no Grupo de Contrato.",
    "Os campos financeiros consolidados (Valor SaaS, Valor Serviço, MRR) estão atualizados no Grupo de Contrato.",
    "Grupos com status = 'Cancelado' não têm itens de linha recriados.",
]:
    add_bullet(doc, item)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 4. VISÃO TÉCNICA
# ════════════════════════════════════════════════════════
add_heading(doc, "4. Visão Técnica", level=1)

# 4.1 Objetos HubSpot
add_heading(doc, "4.1 Objetos HubSpot Envolvidos", level=2)
t41 = create_brand_table(doc,
    ["Objeto", "objectTypeId", "Papel no processo"],
    [
        ["Grupo de Contrato", "2-54707985", "Objeto central. Recebe os valores Oracle (prog_qtd_*) e tem seus itens de linha reconstruídos."],
        ["Contrato", "2-54707915", "Propaga prog_qtd_* para o GC quando um novo contrato é criado."],
        ["Line Items (Itens de Linha)", "0-8 (line_items)", "São arquivados e recriados a cada execução do workflow de split."],
        ["Empresa (Company)", "0-2", "Associada ao GC. O [DEV] Integração Oracle em Company é um workflow auxiliar de desenvolvimento."],
        ["Portfólio", "2-54708014", "Recebe atualização de itens de linha via workflow separado após alterações no GC."],
    ],
)
set_col_widths(t41, [4.5, 3.5, 9.5])
doc.add_paragraph()

# 4.2 Propriedades Críticas
add_heading(doc, "4.2 Propriedades Críticas", level=2)

add_para(doc, "Propriedades de Integração Oracle (grupo: integração_oracle)", bold=True, color_key="navy")
t42a = create_brand_table(doc,
    ["Propriedade HubSpot", "Tipo", "Origem / Finalidade"],
    [
        ["nr_contrato_hs", "string", "Número do contrato Oracle. Chave de busca do GC no Oracle."],
        ["sistema", "enumeration (checkbox)", "Módulos/sistemas Sienge contratados. Determina o split de SaaS/MM."],
        ["prog_qtd_saas", "number", "Valor mensal SaaS vindo do Oracle (VS_LUS.00013)"],
        ["prog_qtd_manutencao", "number", "Valor mensal de Manutenção LU vindo do Oracle (VS_MAN.00122)"],
        ["prog_qtd_apis", "number", "Valor mensal de APIs (VS_LUS.00031)"],
        ["prog_qtd_connectors", "number", "Valor mensal de Connectors (VS_LUS.00032)"],
        ["prog_qtd_consultoria", "number", "Valor de Consultoria (VS_CNS.00001)"],
        ["prog_qtd_locacao_de_data_center", "number", "Data Center (VS_LDT.00009)"],
        ["prog_qtd_base_de_testes", "number", "Base de Testes (VS_LUS.00010)"],
        ["prog_qtd_ecustos_integracao", "number", "E-Custos Integração (VS_LUS.00033)"],
        ["prog_qtd_locacao", "number", "Locação (VS_LCC.00008)"],
        ["prog_qtd_nota_fiscal_eletronica", "number", "NFe (VS_SD.00100)"],
        ["prog_qtd_nota_fiscal_de_servico", "number", "NFs (VS_SD.00094)"],
        ["prog_qtd_emissao_de_conhecimento_de_transporte_eletronico_cte", "number", "CTe (VS_SD.00096)"],
        ["prog_qtd_nota_fiscal_eletronica_projeto_recepcao_nota_fiscal", "number", "Recepção NFe (VS_SD.00071)"],
        ["prog_qtd_intermediacao", "number", "Intermediação (VS_INT.00089)"],
        ["prog_qtd_desenvolvimento_fabrica_de_desenvolvimento", "number", "Dev/Fábrica (VS_DSV.00003)"],
        ["prog_qtd_licenca_de_uso_base_client_share", "number", "Licença de Uso Base Client Share (VS_LCC.00008)"],
        ["last_sync_at_hs", "datetime", "Data/hora da última execução de sync"],
        ["sync_status_hs", "string", "Log de erro da última sync (quando houver)"],
    ],
)
set_col_widths(t42a, [7.5, 2.5, 7.5])
doc.add_paragraph()

add_para(doc, "Propriedades de controle e resultado (grupo: grupos_de_contrato_information)", bold=True, color_key="navy")
t42b = create_brand_table(doc,
    ["Propriedade HubSpot", "Tipo", "Finalidade"],
    [
        ["mrr_oficial", "number", "MRR calculado do GC — alimentado por workflow secundário"],
        ["valor_saas_oficial", "number", "Valor SaaS total — alimentado por workflow secundário"],
        ["valor_servico__oficial", "number", "Valor Serviço total — alimentado por workflow secundário"],
        ["valor_data_center_lu", "number", "Valor Data Center LU — alimentado por workflow secundário"],
        ["valor_manutencao_mensal_lu", "number", "Valor Manutenção Mensal LU — alimentado por workflow secundário"],
        ["gatilho_de_atualizacao_de_itens_de_linha", "bool", "Dispara reconstrução de itens de linha quando um Contrato novo é associado ao GC"],
        ["status", "enumeration", "Ativo ou Cancelado. GCs Cancelados são suprimidos da integração Oracle."],
    ],
)
set_col_widths(t42b, [7.5, 2.5, 7.5])
doc.add_paragraph()

# 4.3 Workflows
add_heading(doc, "4.3 Workflows Envolvidos", level=2)

add_workflow_block(
    doc,
    nome="[DEV] Integração Oracle — Grupo de Contrato",
    id_wf="1758086000",
    objeto="Grupo de Contrato (2-54707985)",
    disparo="Agendado diariamente às 03:30h (schedule). Re-enrollment ativo.",
    acoes=(
        "1. Autentica no Oracle via POST /auth/login com ORACLE_USERNAME / ORACLE_PASSWORD → obtém oracle_token.\n"
        "2. Branch SUCCESS.\n"
        "3. Consulta GET /contratos?p_data={DDMMYYYY}&x_api_token={token}&p_operacao=1 → retorna dados de todos os contratos do dia.\n"
        "4. Branch SUCCESS.\n"
        "5. Processa o JSON retornado: para cada cliente, extrai nr_contrato e calcula grupo_de_contrato_id via regex ^[^._]+ "
        "(prefixo antes do primeiro ponto, underline ou hífen). Acumula total de prog_qtd por grupo."
    ),
    dependencias=(
        "Enrollment atual: um único GC específico (hs_object_id 45054877487) — configuração de desenvolvimento. "
        "O mecanismo de atualização em massa das propriedades prog_qtd_* nos GCs de produção não está mapeado no HubSpot via MCP."
    ),
    status="Ativo (DEV)"
)

add_workflow_block(
    doc,
    nome="[01. Integração Oracle - Grupo de Contrato] Split de serviços em itens de linha (v03)",
    id_wf="1785580697",
    objeto="Grupo de Contrato (2-54707985)",
    disparo=(
        "Re-enrollment ativo. Disparado por qualquer alteração em qualquer propriedade prog_qtd_* ou na propriedade sistema. "
        "Suprimido para GCs com status = 'Cancelado'."
    ),
    acoes=(
        "1. Random delay 1-19 segundos (anti-concorrência inicial).\n"
        "2. A/B Test em 3 branches: direto / wait 1 min / wait 2 min (distribui execuções concorrentes no tempo).\n"
        "3. Batch archive de todos os itens de linha existentes do GC (delete-all).\n"
        "4. Branch SUCCESS do archive.\n"
        "5. Random delay 1-19 segundos.\n"
        "6. Branch 'Com Split' (se prog_qtd_saas OR prog_qtd_manutencao OR prog_qtd_connectors IS_KNOWN):\n"
        "   — Busca produtos HubSpot pelos nomes dos sistemas (campo sistema do GC);\n"
        "   — Split de sistemas: distribui o valor SaaS/MM proporcionalmente pelo hs_price_brl de cada módulo;\n"
        "   — Serviços fixos (APIs, NFe, NFs, CTe, Base Teste, E-Custos, Data Center): preço = prog_qtd_* correspondente;\n"
        "   — Split de conectores: distribui prog_qtd_connectors proporcionalmente entre produtos 'Conector*';\n"
        "   — Contrato misto SaaS+MM: itens de sistema (valor SaaS) + item 'Manutenção Mensal' separado.\n"
        "   — Sem sistemas: item único 'SaaS' ou 'Manutenção' com valor total.\n"
        "7. Branch 'Sem Split' (demais casos):\n"
        "   — Cria item por prog_qtd_* não nulo, usando lin_servico_desc_* como nome e prog_qtd como preço."
    ),
    dependencias=(
        "Requer que sistema esteja preenchido para split proporcional. "
        "Requer produtos com hs_price_brl preenchido no catálogo HubSpot para split funcionar. "
        "Secret: HUBSPOT_INTEGRATIONS."
    )
)

add_workflow_block(
    doc,
    nome="[02.03. Contratos] Novo contrato criado → Cria ou Associa à um Grupo de Contrato + Replica Itens de Linha",
    id_wf="1744643667",
    objeto="Contrato (2-54707915)",
    disparo=(
        "Criação de Contrato com tipo_de_orcamento = 'Primeira Venda', 'Aditivo de Retração' ou 'Aditivo de Expansão'. "
        "Re-enrollment ativo."
    ),
    acoes=(
        "Para Primeira Venda e Expansão com Nova Base / Migração LU→SaaS:\n"
        "  — Cria novo Grupo de Contrato com ~30 propriedades copiadas do Contrato (incluindo todas as prog_qtd_*).\n"
        "Para Aditivo de Expansão — Base Teste:\n"
        "  — Cria GC marcado como Base de Teste.\n"
        "Para Retração / Expansão com GC existente:\n"
        "  — Associa o Contrato ao GC já existente.\n"
        "Em todos os casos:\n"
        "  — Replica itens de linha elegíveis do Contrato para o GC (excluindo: aditivos de usuário, "
        "retração de usuário, Data Center por Usuário, implantação, Licença de Uso LU).\n"
        "  — Para expansão: soma o novo valor ao item existente no GC.\n"
        "  — Para retração: subtrai o valor do item existente no GC."
    ),
    dependencias="Depende da existência do GC para casos de retração/expansão. Secret: nx_interno."
)

add_workflow_block(
    doc,
    nome="[04.04.01] Valor Serviço / [04.04.02] Valor SaaS / [04.04.03] MM LU / [04.04.04] Data Center LU",
    id_wf="1789582978 / 1789452035 / 1789589316 / 1789539100",
    objeto="Grupo de Contrato (2-54707985)",
    disparo="Atualização em prog_qtd_*. Re-enrollment ativo.",
    acoes=(
        "Cada workflow recalcula um campo financeiro consolidado do GC com base nas propriedades prog_qtd_* relevantes:\n"
        "  — 04.04.01: Valor Serviço (valor_servico__oficial)\n"
        "  — 04.04.02: Valor SaaS (valor_saas_oficial)\n"
        "  — 04.04.03: Manutenção Mensal LU (valor_manutencao_mensal_lu)\n"
        "  — 04.04.04: Data Center LU (valor_data_center_lu)"
    ),
    dependencias="Executam em cascata após atualização dos prog_qtd_*."
)

# 4.4 Integração Oracle
add_heading(doc, "4.4 Integração Oracle", level=2)

add_para(doc, "Endpoints:")
t44a = create_brand_table(doc,
    ["Endpoint", "Método", "Finalidade"],
    [
        ["https://api-apex.softplan.com.br/ordsprd/apex_starian/hubspot/auth/login", "POST",
         "Autenticação. Body: {username, password}. Retorna: {token}."],
        ["https://api-apex.softplan.com.br/ordsprd/apex_starian/hubspot/contratos?p_data={DDMMYYYY}&x_api_token={token}&p_operacao=1",
         "GET",
         "Dados de contratos do dia. Retorna: {status_code, l_return, empresas: {clientes: [{contratos: [{nr_contrato, cnpj, itens: [{lin_servico, lin_servico_desc, prog_qtd}]}]}]}}."],
    ],
)
set_col_widths(t44a, [8, 1.5, 8])
doc.add_paragraph()

add_para(doc, "Tabela de mapeamento Oracle → HubSpot (lin_servico → propriedade):")
t44b = create_brand_table(doc,
    ["Código Oracle", "lin_servico_desc Oracle", "Propriedade HubSpot"],
    [
        ["VS_LUS.00013", "SAAS", "prog_qtd_saas"],
        ["VS_MAN.00122", "MANUTENCAO", "prog_qtd_manutencao"],
        ["VS_LUS.00031", "APIs", "prog_qtd_apis"],
        ["VS_LUS.00032", "CONNECTORS", "prog_qtd_connectors"],
        ["VS_CNS.00001", "CONSULTORIA", "prog_qtd_consultoria"],
        ["VS_LDT.00009", "LOCACAO DE DATA CENTER", "prog_qtd_locacao_de_data_center"],
        ["VS_LUS.00010", "BASE DE TESTES", "prog_qtd_base_de_testes"],
        ["VS_LUS.00033", "ECUSTOS INTEGRAÇÃO", "prog_qtd_ecustos_integracao"],
        ["VS_LCC.00008", "LOCACAO", "prog_qtd_locacao"],
        ["VS_SD.00100", "NOTA FISCAL ELETRONICA", "prog_qtd_nota_fiscal_eletronica"],
        ["VS_SD.00094", "NOTA FISCAL DE SERVICO", "prog_qtd_nota_fiscal_de_servico"],
        ["VS_SD.00096", "EMISSAO DE CONHECIMENTO DE TRANSPORTE ELETRONICO (CTE)", "prog_qtd_emissao_de_conhecimento_de_transporte_eletronico_cte"],
        ["VS_SD.00071", "NOTA FISCAL ELETRONICA - PROJETO RECEPCAO NOTA FISCAL", "prog_qtd_nota_fiscal_eletronica_projeto_recepcao_nota_fiscal"],
        ["VS_INT.00089", "INTERMEDIACAO", "prog_qtd_intermediacao"],
        ["VS_DSV.00003", "DESENVOLVIMENTO (FABRICA DE DESENVOLVIMENTO)", "prog_qtd_desenvolvimento_fabrica_de_desenvolvimento"],
        ["VS_LCC.00008 (var)", "LICENCA DE USO BASE (CLIENT SHARE)", "prog_qtd_licenca_de_uso_base_client_share"],
    ],
)
set_col_widths(t44b, [4, 7, 6.5])
doc.add_paragraph()

# 4.5 Cards e Customizações
add_heading(doc, "4.5 Cards e Customizações", level=2)
add_para(doc,
    "Todos os custom codes do processo são ações Python (PYTHON39) ou JavaScript (NODE20X) executadas dentro de workflows HubSpot. "
    "Não há UI Extension ou card customizado específico documentado para o processo 8.1. "
    "Os dados Oracle são visualizados através das propriedades padrão do objeto Grupo de Contrato na interface do HubSpot."
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 5. RISCOS E DEPENDÊNCIAS
# ════════════════════════════════════════════════════════
add_heading(doc, "5. Riscos e Dependências", level=1)

t5 = create_brand_table(doc,
    ["Risco / Dependência", "Impacto", "Mitigação atual"],
    [
        ["Falha na autenticação Oracle (credenciais expiradas ou API indisponível)",
         "Todos os GCs param de ser atualizados. prog_qtd_* ficam desatualizados silenciosamente.",
         "Workflow retorna erro no log. sync_status_hs pode registrar o erro. Não há alerta automático documentado."],
        ["Propriedade sistema vazia para cliente SaaS",
         "Cria item único 'SaaS' com valor total em vez de split por módulo. MRR por módulo fica sem detalhamento.",
         "Regra implementada no workflow (Regra 5). Não é erro — é comportamento esperado para GCs sem módulos mapeados."],
        ["hs_price_brl zerado ou ausente em algum módulo do catálogo",
         "Peso desse módulo no split = 0. Módulo recebe item com preço R$0,00.",
         "Não há validação prévia. Ponto a validar com equipe de produtos."],
        ["Execuções concorrentes do workflow de split",
         "Itens de linha duplicados no GC.",
         "Random delay + A/B test de 3 branches. Eficácia depende do volume de alterações simultâneas."],
        ["Ticket 86ba8pa3r em andamento (Valor Final Serviços)",
         "Campo valor_servico__oficial pode estar desatualizado até o ticket ser resolvido.",
         "Monitorar resolução do ticket."],
        ["Oracle retorna valor divergente do contratado",
         "GC mostra MRR diferente do valor comercialmente acordado.",
         "Alinhado com Sienge que é aceitável — Oracle reflete faturamento real, não valor negociado."],
        ["Workflow [DEV] Integração Oracle em GC apenas com um GC específico no enrollment",
         "O mecanismo real de atualização em massa das propriedades prog_qtd_* nos GCs de produção não está totalmente documentado.",
         "Ponto a validar. Pode haver middleware externo (não HubSpot) responsável pela sync em produção."],
    ],
)
set_col_widths(t5, [5.5, 5, 7])
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 6. MATERIAIS DE APOIO
# ════════════════════════════════════════════════════════
add_heading(doc, "6. Materiais de Apoio", level=1)

t6mat = create_brand_table(doc,
    ["Tipo", "Título / ID", "Localização"],
    [
        ["Documento de orientação", "[Nexforce & Sienge] Orientações integração Oracle → HubSpot", "Drive do projeto / pasta drive do processo 8.1"],
        ["Escopo técnico", "[Sienge] Escopo: Criação de itens de linha para Grupo de Contrato (Custom code)", "Drive do projeto / pasta drive do processo 8.1"],
        ["Ticket ClickUp", "86ba1rdmu — Oracle CTe não puxado", "ClickUp — projeto Sienge RaaS (done)"],
        ["Ticket ClickUp", "86ba1ch4k — GC com valor divergente na barra lateral", "ClickUp — projeto Sienge RaaS (done)"],
        ["Ticket ClickUp", "86ba1exe3 — Segmentar valores por serviços em Deal, Contrato e GC", "ClickUp — projeto Sienge RaaS (done)"],
        ["Ticket ClickUp", "86ba8pa3r — Valor Final Serviços não preenchida no Negócio", "ClickUp — projeto Sienge RaaS (in progress)"],
        ["Workflow HubSpot", "Split de serviços em itens de linha (v03)", "HubSpot Portal 50102745 — id 1785580697"],
        ["Workflow HubSpot", "[DEV] Integração Oracle — Grupo de Contrato", "HubSpot Portal 50102745 — id 1758086000"],
    ],
)
set_col_widths(t6mat, [3, 8, 6.5])
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 7. PONTOS A VALIDAR
# ════════════════════════════════════════════════════════
add_heading(doc, "7. Pontos a Validar", level=1)

add_callout(doc, "warning",
    "Ação requerida — Vinicius Vanoni",
    "Os pontos abaixo não foram confirmados durante a documentação e requerem validação do responsável "
    "(Vinicius Vanoni) ou da equipe técnica antes da publicação definitiva deste documento."
)
doc.add_paragraph()

t7 = create_brand_table(doc,
    ["ID", "Ponto", "Detalhe"],
    [
        ["PV-01", "Mecanismo real de sync em produção",
         "O workflow [DEV] Integração Oracle (id 1758086000) está com enrollment em um único GC específico (id 45054877487). "
         "Como ocorre a atualização das propriedades prog_qtd_* nos demais Grupos de Contrato em produção? "
         "Existe middleware externo? Há outro workflow não mapeado?"],
        ["PV-02", "Valores exatos do campo sistema",
         "O campo sistema é um multiple checkbox no GC. Os valores cadastrados correspondem exatamente aos nomes dos produtos "
         "no catálogo HubSpot (ex: 'Contas a Pagar', 'Vendas')? "
         "Uma divergência de nome impede o split proporcional de funcionar."],
        ["PV-03", "Preços (hs_price_brl) dos módulos no catálogo HubSpot",
         "O split de SaaS/MM usa hs_price_brl de cada módulo como peso. Se o preço estiver zerado ou ausente, "
         "o módulo recebe item com R$0,00. Validar se todos os módulos têm preços cadastrados corretamente."],
        ["PV-04", "Ticket 86ba8pa3r — Valor Final Serviços",
         "Este ticket de estabilização está in progress. Qual é a resolução esperada e quando impacta o processo 8.1?"],
        ["PV-05", "Alertas de falha na sync Oracle",
         "Quando a autenticação Oracle falha ou a API retorna erro, existe algum alerta automático (e-mail, Slack, HubSpot notification)? "
         "Ou a detecção depende de monitoramento manual do campo sync_status_hs?"],
    ],
)
set_col_widths(t7, [1.5, 4.5, 11.5])
doc.add_paragraph()

doc.save(OUTPUT_PATH)
print(f"Documento gerado: {OUTPUT_PATH}")
