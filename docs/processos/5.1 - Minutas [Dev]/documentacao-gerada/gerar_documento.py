"""
Gerador de documentação para o processo 5.1 - Minutas [Dev]
Versão: 1.0
Executar: python3 gerar_documento.py
"""

import sys
from pathlib import Path

# ── nexforce-brand helpers ──────────────────────────────────────────────────
SKILL_DIR = Path("/home/hugo-zanni/Nexforce/Projects/handover-sienge/.claude/skills/nexforce-brand")
sys.path.insert(0, str(SKILL_DIR / "references"))

import python_docx_helpers as _brand
_brand.nexforce_assets = lambda: SKILL_DIR / "assets"

from python_docx_helpers import (
    apply_nexforce_theme,
    insert_brand_header,
    insert_brand_footer,
    create_brand_table,
    add_callout,
    NEXFORCE_COLORS,
    logo_path,
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "5.1 - Minutas [Dev].docx"


def _rgb(key: str) -> RGBColor:
    return RGBColor.from_string(NEXFORCE_COLORS[key])


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Lato"
    if level == 1:
        run.font.color.rgb = _rgb("primary_black")
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = _rgb("navy")
        run.font.size = Pt(13)
        run.font.bold = True
    else:
        run.font.color.rgb = _rgb("dark_gray")
        run.font.size = Pt(11)
        run.font.bold = True
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color_key="dark_gray", indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Lato"
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color_key)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    for run in p.runs:
        run.font.name = "Lato"
        run.font.size = Pt(10.5)
        run.font.color.rgb = _rgb("dark_gray")
    return p


def add_workflow_block(doc, nome, id_wf, objetivo, objeto, disparo, acoes, dependencias=None, status="Ativo"):
    p = doc.add_paragraph()
    run = p.add_run(f"Workflow: {nome}")
    run.bold = True
    run.font.name = "Lato"
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb("navy")

    rows_data = [
        ("ID", id_wf),
        ("Objetivo", objetivo),
        ("Status", status),
        ("Objeto HubSpot", objeto),
        ("Disparo", disparo),
        ("Ações principais", acoes),
    ]
    if dependencias:
        rows_data.append(("Dependências", dependencias))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        cell_l = table.rows[i].cells[0]
        cell_l.text = label
        _set_cell_bg(cell_l, NEXFORCE_COLORS["callout_info_bg"])
        for para in cell_l.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Lato"
                run.font.size = Pt(9.5)
                run.font.color.rgb = _rgb("navy")

        cell_v = table.rows[i].cells[1]
        cell_v.text = value
        bg = NEXFORCE_COLORS["white"] if i % 2 == 0 else NEXFORCE_COLORS["near_white"]
        _set_cell_bg(cell_v, bg)
        for para in cell_v.paragraphs:
            for run in para.runs:
                run.font.name = "Lato"
                run.font.size = Pt(9.5)
                run.font.color.rgb = _rgb("dark_gray")

    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(13)
    doc.add_paragraph()


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
apply_nexforce_theme(doc)

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

insert_brand_header(doc, title="5.1 — Minutas [Dev]", period="Junho 2026")
insert_brand_footer(doc, confidentiality="Uso interno  ·  Nexforce Services  ·  v1.0")

# ════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Cm(8)
r = p_sub.add_run("Documentação de Processos — HubSpot RaaS")
r.font.name = "Lato"
r.font.size = Pt(12)
r.font.color.rgb = _rgb("mid_gray")

doc.add_paragraph()

p_title = doc.add_heading("5.1 — Minutas [Dev]", level=0)
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in p_title.runs:
    run.font.name = "Lato"
    run.font.color.rgb = _rgb("primary_black")
    run.font.size = Pt(26)

doc.add_paragraph()

meta_table = create_brand_table(
    doc,
    headers=["Campo", "Valor"],
    rows=[
        ["Cliente", "Sienge"],
        ["Responsável pelo processo", "Jorge Souza"],
        ["Empresa responsável", "Nexforce Services"],
        ["Data de documentação", "16 de junho de 2026"],
        ["Versão", "1.0"],
        ["Status do processo", "Estabilização pós Go Live"],
    ],
)
set_col_widths(meta_table, [6, 10])

doc.add_page_break()

# ════════════════════════════════════════════════════════
# SUMÁRIO
# ════════════════════════════════════════════════════════
add_heading(doc, "Sumário", level=1)
for item in [
    "1. Objetivo",
    "2. Contexto de Negócio",
    "3. Visão Funcional",
    "   3.1 Fluxo Operacional",
    "   3.2 Regras de Negócio",
    "   3.3 Critério de Validação",
    "4. Visão Técnica",
    "   4.1 Objetos HubSpot Envolvidos",
    "   4.2 Propriedades Críticas",
    "   4.3 Workflows Envolvidos",
    "   4.4 Templates e Módulos HubL",
    "   4.5 HubDB — Divisão de Faturamento",
    "   4.6 Cards e Customizações",
    "5. Riscos e Dependências",
    "6. Materiais de Apoio",
    "7. Pontos a Validar",
]:
    add_para(doc, item, size=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. OBJETIVO
# ════════════════════════════════════════════════════════
add_heading(doc, "1. Objetivo", level=1)
add_para(doc,
    "Desenvolver e manter os templates de orçamento personalizado (minutas contratuais) no HubSpot Design Manager, "
    "garantindo que cada tipo de operação comercial — primeira venda, upsell, downsell, cessão de direitos, troca de "
    "modalidade, base teste e retração — tenha um template HubL correspondente que puxe automaticamente as propriedades "
    "do Negócio, Empresa e itens de linha ao gerar o orçamento no HubSpot."
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 2. CONTEXTO DE NEGÓCIO
# ════════════════════════════════════════════════════════
add_heading(doc, "2. Contexto de Negócio", level=1)

add_para(doc,
    "O Sienge opera com contratos comerciais de alta complexidade: múltiplas modalidades de aquisição (SaaS e LU), "
    "dezenas de módulos de software contratáveis, anuentes (empresas do grupo que compartilham o contrato), "
    "divisão de faturamento entre CNPJs e tipos variados de aditivos."
)
doc.add_paragraph()
add_para(doc,
    "Para formalizar cada operação, o time comercial gera um orçamento (Quote) no HubSpot associado ao Negócio. "
    "Esse orçamento é o documento contratual enviado ao cliente para assinatura digital. O conteúdo do documento "
    "— cláusulas, valores, partes, quadros de assinatura, divisão de faturamento — é gerado dinamicamente a partir "
    "de um template HubL configurado no HubSpot Design Manager."
)
doc.add_paragraph()
add_para(doc,
    "O processo 5.1 cobre o desenvolvimento técnico desses templates: criação, manutenção e correção dos módulos "
    "HubL que compõem cada tipo de minuta, bem como a integração com o HubDB que armazena a divisão de faturamento "
    "entre canais/anuentes."
)
doc.add_paragraph()

add_para(doc, "Por que existe separação entre 5.0 e 5.1?", bold=True, color_key="primary_black")
add_para(doc,
    "O processo 5.0 (Minutas) cobre a operação funcional: qual template selecionar, como preencher os campos, "
    "fluxo de aprovação e assinatura. O processo 5.1 cobre a camada técnica: como os templates são construídos "
    "em HubL, quais módulos existem, quais propriedades são lidas e como a lógica condicional funciona dentro do "
    "Design Manager.",
    size=10.5, indent=True
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 3. VISÃO FUNCIONAL
# ════════════════════════════════════════════════════════
add_heading(doc, "3. Visão Funcional", level=1)

# 3.1 Fluxo Operacional
add_heading(doc, "3.1 Fluxo Operacional", level=2)
add_para(doc, "Do ponto de vista funcional, o processo de geração de minuta segue os passos abaixo:")
doc.add_paragraph()

steps = [
    ("1", "Vendedor identifica o tipo de operação",
     "O tipo de orçamento (Primeira Venda, Upsell, Downsell, Cessão de Direitos, Troca de Modalidade, Base Teste etc.) "
     "é selecionado pelo vendedor na propriedade 'Tipo de Orçamento' do Negócio."),
    ("2", "Workflow copia propriedades da Empresa para o Negócio",
     "O workflow [Minutas] 01 (id 1793711661) é disparado na criação do Negócio e copia as propriedades "
     "da Empresa (CNPJ, razão social, endereço, representante legal) para o Negócio, garantindo que o template HubL "
     "acesse esses dados diretamente no objeto Deal."),
    ("3", "Negócio avança para etapa Contrato",
     "O workflow [Minutas] 02 (id 1793738566) é disparado quando o Negócio entra na etapa 'Contrato' "
     "e copia propriedades adicionais do Negócio de volta para a Empresa (número de contrato gerado, data de assinatura, etc.)."),
    ("4", "Vendedor seleciona o template e gera o Quote",
     "O vendedor acessa o objeto Negócio, clica em 'Orçamento', seleciona o template correspondente ao tipo de operação "
     "e clica em 'Publicar / Assinar e Enviar'. O HubSpot carrega o template HubL escolhido e injeta as propriedades do "
     "Negócio, Empresa e itens de linha em tempo real."),
    ("5", "Fluxo de aprovação (quando aplicável)",
     "Para orçamentos que requerem aprovação, o workflow nativo 'Quando os orçamentos exigem aprovação' (id 1741583212) "
     "dispara e roteia para os aprovadores definidos no processo 3.0 (Aprovações)."),
    ("6", "Orçamento publicado e enviado para assinatura",
     "Após aprovação (ou diretamente, se aprovação não for exigida), o workflow 'Quando os orçamentos são aprovados' "
     "(id 1741583214) e os workflows Nx. [Alerta de assinaturas] (ids 1790743868/1790744223/1790745492) notificam "
     "as partes e registram o status da assinatura."),
    ("7", "Assinaturas completadas",
     "O workflow Nx. 01.02. [Alerta de assinaturas] → Assinaturas completadas (id 1790745492) dispara ao completar "
     "todas as assinaturas digitais. Para fluxos de Expansão, o negócio avança para 'Upsell'. "
     "Para Retração, o ticket de CS é movido de 'Formalização' para 'Retração Realizada'."),
]

for num, titulo, descricao in steps:
    p = doc.add_paragraph()
    r_num = p.add_run(f"Passo {num}: ")
    r_num.bold = True
    r_num.font.name = "Lato"
    r_num.font.color.rgb = _rgb("navy")
    r_num.font.size = Pt(11)
    r_titulo = p.add_run(titulo)
    r_titulo.bold = True
    r_titulo.font.name = "Lato"
    r_titulo.font.size = Pt(11)
    r_titulo.font.color.rgb = _rgb("primary_black")
    add_para(doc, descricao, size=10.5, indent=True)
    doc.add_paragraph()

# 3.2 Regras de Negócio
add_heading(doc, "3.2 Regras de Negócio", level=2)

regras = [
    ("Regra 1 — Um template por tipo de operação e modalidade",
     "Cada combinação de tipo de operação + modalidade (SaaS ou LU) tem um template dedicado. "
     "Templates SaaS e LU diferem no conteúdo de cláusulas de valor, na forma de calcular MRR e nas partes "
     "signatárias. Não existe template genérico que sirva para ambos."),
    ("Regra 2 — Propriedade 'Número do Contrato Gerado' é a chave das minutas",
     "As minutas de Cessão de Direitos, Upsell API e demais tipos de aditivo devem puxar o número do contrato da "
     "propriedade 'Número do Contrato Gerado' (gerada automaticamente pelo workflow 1795328011), não da propriedade "
     "'Número do Contrato' (que pode vir no formato 'nº+traço+aditivo' — inadequado para uso em cláusulas)."),
    ("Regra 3 — Divisão de faturamento lida do HubDB",
     "O quadro de divisão de faturamento (canal/anuente, percentual, CNPJ) é armazenado em uma tabela HubDB. "
     "O template HubL lê essa tabela via API HubDB no momento da geração do orçamento. "
     "Linhas duplicadas no HubDB (status 'em espera') causam divisão duplicada na minuta — era necessário deletar "
     "manualmente as linhas duplicadas como paliativo durante o período de estabilização."),
    ("Regra 4 — Implantação não entra na divisão de faturamento nem no valor do negócio",
     "Itens de linha do tipo 'Implantação' não devem ser somados ao valor total do negócio, não entram na trava "
     "de divisão de faturamento e não aparecem na minuta. Apenas o canal Starian lança contratos com implantação — "
     "e para eles o fluxo de aprovação do financeiro não se aplica."),
    ("Regra 5 — Templates de CS (retração) requerem signatários específicos",
     "As minutas de CS para o produto Sienge Plataforma devem incluir Maria Madalena e Giovani como signatários. "
     "Além de representantes legais, os templates devem suportar 'testemunhas' (podendo ser apenas visualizadores "
     "ou também assinantes)."),
    ("Regra 6 — Lógica condicional por tipo de serviço nas minutas de NF",
     "Na minuta de consultoria/implantação: quando o orçamento tiver item 'NFs', puxar 'Recepção/Emissão de NFSE'; "
     "quando tiver 'NFe', puxar 'Recepção/Emissão de NFE'. Outros serviços (Ecustos, CTe) não aparecem. "
     "Na minuta de NFs: o rótulo do serviço NFe deve ser 'Nota Fiscal Eletrônica' (sem a palavra 'Emissão')."),
    ("Regra 7 — Valor na minuta deve ser o MRR resultante, não o delta",
     "Nas minutas de downsell/retração, o valor exibido é o novo MRR do cliente após a redução "
     "(valor atual menos o delta reduzido), não o delta isolado. "
     "Exemplo: MRR atual R$ 6.588,32, redução R$ 169,59 → minuta mostra R$ 6.418,73."),
    ("Regra 8 — Base Teste: novo nº de contrato e valor da cláusula 2.1.2",
     "Minutas de Base Teste devem gerar um novo número de contrato (não usar o nº do contrato pai). "
     "A cláusula 2.1.2 usa o valor de R$ 600,00 (atualizado de R$ 390,00 via ticket de estabilização). "
     "A quantidade de usuários na minuta de Base Teste é tabelada com base no MRR do cliente."),
    ("Regra 9 — Troca de modalidade: nº do contrato e divisão de faturamento",
     "Minutas de Troca de Modalidade, Expansão de Nova Base e Migração GO→Plataforma devem gerar novo nº de contrato. "
     "A divisão de faturamento deve mostrar apenas a linha correspondente ao escopo (ex: API), não toda a divisão do cliente."),
    ("Regra 10 — Geração automática de Quote substituída por fluxo manual",
     "O workflow de geração automática de orçamento para retração (id 1793114638) está desativado e marcado como "
     "'desatualizado'. O fluxo atual exige que o CS selecione manualmente o modelo de minuta e clique em 'Assinar e Enviar'. "
     "O workflow Nx. [Fluxo de envio para ganho manual] - Retração (id 1793664676) gerencia esse fluxo."),
]

for titulo, texto in regras:
    add_para(doc, titulo, bold=True, color_key="primary_black")
    add_para(doc, texto, size=10.5, indent=True)
    doc.add_paragraph()

# 3.3 Critério de Validação
add_heading(doc, "3.3 Critério de Validação", level=2)
add_para(doc, "A minuta foi gerada corretamente quando:")
for item in [
    "O template selecionado corresponde ao tipo de operação e à modalidade (SaaS ou LU) do cliente.",
    "O número do contrato exibido na minuta é o 'Número do Contrato Gerado' (sem traço-aditivo).",
    "Os valores monetários na minuta coincidem com os valores dos itens de linha do Negócio.",
    "O quadro de divisão de faturamento reflete exatamente a distribuição cadastrada, sem linhas duplicadas.",
    "O quadro de assinaturas contém todas as partes exigidas pelo tipo de minuta (incluindo testemunhas quando aplicável).",
    "Itens de linha do tipo 'Implantação' não aparecem na minuta nem no valor total.",
    "A quantidade de usuários na minuta coincide com o valor exibido na UI Extensions.",
    "Nas minutas de retração, o valor exibido é o MRR resultante (novo total), não o delta reduzido.",
]:
    add_bullet(doc, item)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 4. VISÃO TÉCNICA
# ════════════════════════════════════════════════════════
add_heading(doc, "4. Visão Técnica", level=1)

# 4.1 Objetos HubSpot
add_heading(doc, "4.1 Objetos HubSpot Envolvidos", level=2)
t41 = create_brand_table(doc,
    ["Objeto", "objectTypeId", "Papel no processo"],
    [
        ["Negócio (Deal)", "0-3", "Objeto principal. Contém tipo de orçamento, valores, itens de linha e propriedades copiadas da Empresa para uso no template."],
        ["Orçamento (Quote)", "0-14", "Objeto onde o template HubL é instanciado. Associado ao Negócio. Gerencia o fluxo de aprovação e assinatura."],
        ["Empresa (Company)", "0-2", "Fornece CNPJ, razão social, endereço e dados do representante legal para o template. Recebe dados de volta após assinatura."],
        ["Contato (Contact)", "0-1", "Signatários e representantes legais associados ao orçamento."],
        ["Contrato (objeto custom)", "2-54707915", "Criado após o ganho do negócio (workflow 1744636053). O template de minuta referencia propriedades do contrato em alguns casos."],
        ["Grupo de Contrato", "2-54707985", "Referenciado em minutas de Cessão de Direitos e Troca de Modalidade para obter o 'Nº Contrato Pai'."],
        ["Requisição", "2-54232449", "Objeto onde o CS registra o pedido de retração. Associado ao Negócio de retração que gera a minuta."],
    ],
)
set_col_widths(t41, [4, 3, 10.5])
doc.add_paragraph()

# 4.2 Propriedades Críticas
add_heading(doc, "4.2 Propriedades Críticas", level=2)

add_para(doc, "Propriedades do Negócio (Deal) usadas pelos templates", bold=True, color_key="navy")
t42a = create_brand_table(doc,
    ["Propriedade", "Tipo", "Uso no template"],
    [
        ["tipo_de_orcamento", "enumeration", "Determina qual template é exibido ao vendedor e qual fluxo de workflow é disparado."],
        ["numero_do_contrato_gerado", "string", "Número do contrato a ser exibido nas cláusulas contratuais da minuta (formato sem traço-aditivo)."],
        ["numero_do_contrato", "string", "Propriedade legada — usada apenas em minutas que não foram migradas para 'Número do Contrato Gerado'."],
        ["numero_do_aditivo", "string", "Número do aditivo. Exibido no cabeçalho de minutas de expansão/retração."],
        ["modalidade_de_aquisicao", "enumeration", "SaaS ou LU — determina qual variante do template é carregada (cláusulas de valor diferem)."],
        ["amount", "number", "Valor total do negócio. Exibido como valor da minuta. Implantação não deve ser somada aqui."],
        ["data_de_vencimento_do_contrato", "date", "Data de vencimento exibida nas cláusulas de vigência."],
        ["indice_de_reajuste", "enumeration", "Índice de reajuste anual (IPCA, IGP-M, etc.) exibido nas cláusulas contratuais."],
        ["cnpj_do_cliente", "string", "CNPJ exibido no cabeçalho da minuta (copiado da Empresa pelo workflow [Minutas] 01)."],
        ["razao_social", "string", "Razão social exibida no cabeçalho da minuta."],
        ["nome_do_representante_legal", "string", "Nome do representante legal para o quadro de assinaturas."],
        ["cpf_representante_legal", "string", "CPF do representante legal. Necessário para minutas com CPF (ex: Redução de Escopo SaaS)."],
    ],
)
set_col_widths(t42a, [5.5, 2.5, 9.5])
doc.add_paragraph()

add_para(doc, "Propriedades de orçamento (Quote) controladas pelos workflows nativos", bold=True, color_key="navy")
t42b = create_brand_table(doc,
    ["Propriedade", "Tipo", "Uso"],
    [
        ["hs_quote_status", "enumeration", "Status do orçamento: DRAFT, PENDING_APPROVAL, APPROVED, REJECTED, PUBLISHED, SIGNED."],
        ["hs_esign_num_signers_signed", "number", "Número de signatários que já assinaram. Dispara Nx. 01.02. quando igual ao total."],
        ["hs_quote_template_name", "string", "Nome do template HubL selecionado pelo vendedor. Registra qual template foi usado."],
    ],
)
set_col_widths(t42b, [5.5, 2.5, 9.5])
doc.add_paragraph()

# 4.3 Workflows
add_heading(doc, "4.3 Workflows Envolvidos", level=2)

add_workflow_block(
    doc,
    nome="[Minutas] 01. Negócio Criado → Envia Infos da Empresa para Negócio",
    id_wf="1793711661",
    objetivo="Copiar dados cadastrais da Empresa (CNPJ, razão social, endereço, representante legal) para o Negócio no momento da criação, para que o template HubL acesse esses dados diretamente no Deal.",
    objeto="Negócio/Deal (0-3)",
    disparo="Criação de Negócio.",
    acoes="Copia propriedades da Empresa associada para o Negócio: CNPJ, razão social, endereço, nome e CPF do representante legal.",
    dependencias="Empresa deve estar associada ao Negócio no momento da criação. Requer que as propriedades de dados cadastrais estejam preenchidas na Empresa."
)

add_workflow_block(
    doc,
    nome="[Minutas] 02. Etapa = Contrato → Envia Infos do Negócio para Empresa",
    id_wf="1793738566",
    objetivo="Após o Negócio entrar na etapa 'Contrato', copiar propriedades contratuais do Negócio de volta para a Empresa (número de contrato gerado, data de assinatura).",
    objeto="Negócio/Deal (0-3)",
    disparo="Negócio entra na etapa 'Contrato'.",
    acoes="Copia número do contrato gerado e data de vigência do Negócio para a Empresa associada.",
    dependencias="Workflow [01.01. Negócio] Negócio Ganho → Gera Número de Contrato (id 1795328011) deve ter executado antes."
)

add_workflow_block(
    doc,
    nome="Nx. [Fluxo de envio para ganho manual] - Identificação do modelo de orçamento",
    id_wf="1791307187",
    objetivo="Identificar o modelo de orçamento (tipo de minuta) associado ao Quote após o vendedor selecionar 'Assinar e Enviar', para determinar qual fluxo de ganho aplicar.",
    objeto="Orçamento/Quote (0-14)",
    disparo="Mudança de status do orçamento (publicação ou seleção de modelo).",
    acoes="Lê o nome do template HubL selecionado (hs_quote_template_name) e define propriedades de controle que orientam os workflows de ganho do Negócio e de retração.",
    dependencias="Depende do template ter sido selecionado pelo vendedor. Alimenta os workflows Nx. [Fluxo de envio para ganho manual] - Aquisição (1791344691) e Nx. [Fluxo de envio para ganho manual] - Retração (1793664676)."
)

add_workflow_block(
    doc,
    nome="Nx. [Fluxo de envio para ganho manual] - Aquisição",
    id_wf="1791344691",
    objetivo="Processar o ganho de Negócios de Aquisição após a assinatura do orçamento.",
    objeto="Negócio/Deal (0-3)",
    disparo="Orçamento de Aquisição assinado.",
    acoes="Move o Negócio para 'Vendido' e dispara os workflows de criação de Contrato e Onboarding.",
    dependencias="Depende da identificação do modelo pelo workflow 1791307187."
)

add_workflow_block(
    doc,
    nome="Nx. [Fluxo de envio para ganho manual] - Retração",
    id_wf="1793664676",
    objetivo="Processar o ganho de Negócios de Retração após a assinatura do orçamento de CS.",
    objeto="Negócio/Deal (0-3)",
    disparo="Orçamento de Retração assinado.",
    acoes=(
        "Move o Negócio de retração para fechado. "
        "Dispara criação do Contrato de retração. "
        "Atualiza o Ticket de CS para 'Retração Realizada'."
    ),
    dependencias="Depende da identificação do modelo pelo workflow 1791307187. Ticket de CS associado deve existir."
)

add_workflow_block(
    doc,
    nome="Quando os orçamentos exigem aprovação / são aprovados / têm alterações solicitadas",
    id_wf="1741583212 / 1741583214 / 1741583213",
    objetivo="Gerenciar o fluxo de aprovação interna dos orçamentos: notificar aprovadores, processar aprovação/reprovação e liberar o orçamento para assinatura.",
    objeto="Orçamento/Quote (0-14)",
    disparo=(
        "1741583212: orçamento entra em PENDING_APPROVAL.\n"
        "1741583214: orçamento é aprovado (APPROVED).\n"
        "1741583213: alterações são solicitadas (REJECTED)."
    ),
    acoes="Workflows nativos HubSpot. Notificam aprovadores, registram decisão e alteram status do orçamento.",
    dependencias="Regras de aprovação configuradas no processo 3.0 (Aprovações)."
)

add_workflow_block(
    doc,
    nome="Nx. 01.00/01/02. [Alerta de assinaturas]",
    id_wf="1790743868 / 1790744223 / 1790745492",
    objetivo="Notificar o time sobre o progresso das assinaturas digitais: orçamento aprovado, assinaturas iniciadas e assinaturas completadas.",
    objeto="Orçamento/Quote (0-14)",
    disparo=(
        "1790743868: orçamento aprovado.\n"
        "1790744223: primeira assinatura realizada.\n"
        "1790745492: todas as assinaturas completadas."
    ),
    acoes="Enviam notificações internas (e-mail ou HubSpot notification) ao proprietário do Negócio e ao time de contratos.",
    dependencias="Requer que os signatários estejam configurados corretamente no orçamento."
)

doc.add_paragraph()

# 4.4 Templates e Módulos HubL
add_heading(doc, "4.4 Templates e Módulos HubL", level=2)

add_para(doc,
    "Os templates e módulos abaixo foram documentados com base no DOCX original gerado em 22/mai/2026 (responsável: Jorge Souza). "
    "Os nomes exatos no Design Manager não foram confirmados via MCP — os templates listados abaixo são os identificados "
    "nos tickets de estabilização do ClickUp e na documentação original. Ver Pontos a Validar (PV-01).",
    italic=True, size=10, color_key="mid_gray"
)
doc.add_paragraph()

add_para(doc, "Templates de Orçamento Personalizado — identificados nos tickets de estabilização", bold=True, color_key="navy")
t44a = create_brand_table(doc,
    ["Template (nome aproximado)", "Tipo de operação", "Modalidade", "Observações de estabilização"],
    [
        ["Primeira Venda SaaS", "Nova venda", "SaaS", "Correção de bloco duplicado (cláusula 2); ajuste de quantidade de usuários; erro de digitação em cláusula 2.2."],
        ["Primeira Venda GO / LU", "Nova venda", "LU (Gestor Obras)", "Incluir data de treinamento (implantação) que não estava sendo puxada."],
        ["Upsell Usuários e Sistemas SaaS", "Aditivo de expansão", "SaaS", "Divisão de faturamento não deve puxar serviços em upsell de u&s."],
        ["Upsell Usuários e Sistemas LU", "Aditivo de expansão", "LU", "Minutas de Base Teste LU não puxavam usuários (valor tabelado por MRR)."],
        ["Upsell APIs", "Aditivo de expansão", "SaaS/LU", "Remoção de cláusulas indevidas da minuta de venda nova; ajuste de nº do contrato na cláusula 3.1; correção de linhas de anuentes."],
        ["Upsell Serviços NFs", "Aditivo de expansão", "SaaS", "Quadro de divisão de faturamento não estava saindo."],
        ["Downsell Usuários e Sistemas SaaS", "Aditivo de retração", "SaaS", "Configuração para CPF; valor exibido deve ser o MRR resultante (não o delta)."],
        ["Downsell Usuários e Sistemas LU", "Aditivo de retração", "LU", "Correção de valores e de puxada de usuários."],
        ["Downsell Serviços", "Aditivo de retração", "SaaS", "Renomear cabeçalho para incluir nº do aditivo no título."],
        ["Redução Temporária", "Aditivo de retração temporária", "SaaS", "Renomear cabeçalho para incluir nº do aditivo no título."],
        ["Cessão de Direitos", "Cessão", "SaaS/LU", "Nº do contrato deve puxar de 'Número do Contrato Gerado'. Ajuste do Nº Contrato Pai. Puxar nomes dos assinantes do cliente."],
        ["Troca de Modalidade GO→SaaS", "Troca de modalidade", "GO→SaaS", "Divisão de faturamento deve mostrar apenas a linha do escopo (ex: API). Formato de data com ano completo. Geração de novo nº de contrato."],
        ["Troca de Modalidade GO (pagamento)", "Troca de modalidade", "GO", "Template inexistente até estabilização — foi criado durante o processo (ticket 86b9vy8kb)."],
        ["Base Teste SaaS", "Aditivo", "SaaS", "Novo nº de contrato (não usar pai). Cláusula 2.1.2 = R$ 600,00. Divisão de faturamento. Quantidade de usuários tabelada por MRR."],
        ["Base Teste LU (com DC)", "Aditivo", "LU", "Não puxava usuários para clientes LU com DC (MRR Saas = zero)."],
        ["Consultoria / Implantação", "Serviço", "SaaS", "Puxar serviços NFs/NFe condicionalmente. Fluxo de aprovação diferente para canal Starian."],
    ],
)
set_col_widths(t44a, [5.5, 3.5, 2, 6.5])
doc.add_paragraph()

add_para(doc, "Estrutura de módulos HubL", bold=True, color_key="navy")
doc.add_paragraph()

t44b = create_brand_table(doc,
    ["Módulo", "Templates que utilizam"],
    [
        ["Quadro Reativação de Usuários", "Upsell Reativação de Usuários"],
        ["Termos", "Base Teste"],
        ["Assinatura Eletrônica", "Cancelamento Plataforma"],
        ["Considerando Que", "Cancelamento Plataforma"],
        ["Da Resilição", "Cancelamento Plataforma"],
        ["Disposições Gerais", "Cancelamento Plataforma"],
        ["Quadro Cedente", "Cessão de Direitos, Cessão de Direitos GO"],
        ["Quadro Cessionária", "Cessão de Direitos, Cessão de Direitos GO"],
        ["Cabeçalho", "Consultoria Implantação"],
        ["Quadro Preço", "Consultoria Implantação"],
        ["Quadro Serviços Contratados", "Consultoria Implantação"],
        ["Seções", "Consultoria Implantação"],
        ["DOWNSELL SAAS - Quadro Preço", "Downsell Sistemas e/ou Usuários — SaaS"],
        ["DOWSELL LU DC - Quadro Redução de Escopos", "Downsell Usuários — LU com DC"],
        ["Quadro Redução de Serviços", "Downsell Serviços"],
        ["SAAS - Quadro Redução de Escopos", "Downsell Sistemas e/ou Usuários — SaaS"],
        ["Termos", "Downsell Sistemas e/ou Usuários — SaaS"],
        ["Clausulas", "Nova Venda GO"],
        ["Quadro Módulos", "Nova Venda GO"],
        ["Quadro Plano Contratado e Valores", "Nova Venda GO"],
        ["Quadro Planos", "Nova Venda GO"],
        ["ANUENTES - Assinaturas", "Base Teste, Cancelamento Plataforma, Consultoria Implantação, Divisão de Faturamento, Downsell Serviços, Downsell Sistemas e/ou Usuários — SaaS, Nova Venda GO, Primeira Venda, Redução Temporária, Resilição Migração GO para Plataforma, Troca de API, Troca de Modalidade LU com DC + Termo de Resilição, Upsell API, Upsell Conector, Upsell Data Center Exclusivo, Upsell E-Custos, Upsell NFs, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["ANUENTES - Preços e Condições de Pagamento", "Resilição Migração GO para Plataforma"],
        ["Divisão de Faturamento", "Base Teste, Divisão de Faturamento, Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Nova Venda GO, Primeira Venda, Redução Temporária, Resilição Migração GO para Plataforma, Troca de API, Troca de Modalidade LU com DC + Termo de Resilição, Upsell E-Custos, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["Divisão de Faturamento Filtro", "Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Troca de API, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["Quadro Anuentes", "Base Teste, Divisão de Faturamento, Downsell Serviços, Downsell Sistemas e/ou Usuários — SaaS, Nova Venda GO, Primeira Venda, Redução Temporária, Resilição Migração GO para Plataforma, Troca de API, Troca de Modalidade LU com DC + Termo de Resilição, Upsell API, Upsell Conector, Upsell E-Custos, Upsell NFs, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["SERVIÇOS - Divisão de Faturamento", "Upsell API, Upsell Conector, Upsell NFs"],
        ["Assinatura Eletrônica", "Primeira Venda, Resilição Migração GO para Plataforma"],
        ["Condição E-Custos", "Primeira Venda, Resilição Migração GO para Plataforma, Upsell Usuários e Sistemas — LU com DC"],
        ["Disposições Gerais", "Primeira Venda, Resilição Migração GO para Plataforma"],
        ["Foro de Eleição", "Primeira Venda, Resilição Migração GO para Plataforma"],
        ["Quadro Parâmetros Contratação", "Primeira Venda"],
        ["Quadro Preço", "Primeira Venda, Resilição Migração GO para Plataforma"],
        ["Quadro Serviços Conexos", "Primeira Venda, Resilição Migração GO para Plataforma, Troca de Modalidade LU com DC + Termo de Resilição"],
        ["Quadro Serviços Contratados", "Primeira Venda, Resilição Migração GO para Plataforma"],
        ["Quadro Sistemas Licenciados", "Primeira Venda, Resilição Migração GO para Plataforma, Troca de Modalidade LU com DC + Termo de Resilição"],
        ["Preços e Pagamento sem DF", "Resilição Migração GO para Plataforma"],
        ["Quadro Preço", "Redução Temporária"],
        ["Preços e Condições TROCA DM", "Troca de Modalidade LU com DC + Termo de Resilição"],
        ["Preços e Condições TROCA DM Com DF", "Troca de Modalidade LU com DC + Termo de Resilição"],
        ["REATIVAÇÃO USUÁRIOS - Quadro Preço", "Upsell Reativação de Usuários"],
        ["LU Quadro Aumento de Escopos", "Upsell Usuários e Sistemas — LU com DC"],
        ["SAAS Quadro Aumento de Escopos", "Upsell Usuários e Sistemas"],
        ["Termos", "Upsell Usuários e Sistemas"],
        ["Assinaturas", "Base Teste, Cancelamento Plataforma, Cessão de Direitos, Cessão de Direitos GO, Consultoria Implantação, Downsell Serviços, Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Nova Venda GO, Primeira Venda, Redução Temporária, Resilição Migração GO para Plataforma, Troca de API, Troca de Modalidade LU com DC + Termo de Resilição, Upsell API, Upsell Conector, Upsell Data Center Exclusivo, Upsell E-Custos, Upsell NFs, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["Assinaturas - Cessão de direitos", "Cessão de Direitos, Cessão de Direitos GO"],
        ["Cabeçalho", "Base Teste, Cessão de Direitos GO, Divisão de Faturamento, Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Nova Venda GO, Primeira Venda, Redução Temporária, Resilição Migração GO para Plataforma, Troca de Modalidade LU com DC + Termo de Resilição, Upsell Data Center Exclusivo, Upsell E-Custos, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["Cabeçalho Aditivo", "Divisão de Faturamento, Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Redução Temporária, Upsell Data Center Exclusivo, Upsell E-Custos, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["Cidade Estado", "Base Teste, Cessão de Direitos, Cessão de Direitos GO, Consultoria Implantação, Divisão de Faturamento, Downsell Serviços, Downsell Sistemas e/ou Usuários — SaaS, Nova Venda GO, Redução Temporária, Resilição Migração GO para Plataforma, Troca de Modalidade LU com DC + Termo de Resilição, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["QUADRO PREÇO SAAS", "Upsell Usuários e Sistemas"],
        ["Resilição GO - Quadro Parâmetros Contratação", "Resilição Migração GO para Plataforma"],
        ["Sistemas Licenciados BASE TESTE", "Base Teste"],
        ["Tabela Licenciada", "Base Teste, Cancelamento Plataforma, Consultoria Implantação, Divisão de Faturamento, Downsell Serviços, Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Nova Venda GO, Primeira Venda, Redução Temporária, Resilição Migração GO para Plataforma, Troca de API, Troca de Modalidade LU com DC + Termo de Resilição, Upsell API, Upsell Conector, Upsell Data Center Exclusivo, Upsell E-Custos, Upsell NFs, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
        ["Tabela Licenciante", "Todos"],
        ["Versão Sienge", "Base Teste, Troca de Modalidade LU com DC + Termo de Resilição"],
        ["header", "Todos"],
        ["quadro-e-custos", "Downsell Serviços, Downsell Sistemas e/ou Usuários — SaaS, Downsell Usuários — LU com DC, Nova Venda GO, Troca de Modalidade LU com DC + Termo de Resilição, Upsell NFs, Upsell Reativação de Usuários, Upsell Usuários e Sistemas, Upsell Usuários e Sistemas — LU com DC"],
    ],
)
set_col_widths(t44b, [6, 11.5])
doc.add_paragraph()

# 4.5 HubDB
add_heading(doc, "4.5 HubDB — Divisão de Faturamento", level=2)
add_para(doc,
    "A divisão de faturamento entre canais/anuentes é armazenada em uma tabela HubDB no portal do Sienge. "
    "O template HubL realiza uma chamada à API HubDB no momento da geração do orçamento para carregar as linhas "
    "da tabela correspondente ao Negócio atual."
)
doc.add_paragraph()

t45 = create_brand_table(doc,
    ["Dimensão", "Detalhe"],
    [
        ["Estrutura da tabela", "Cada linha representa uma parcela de faturamento: canal/anuente (CNPJ, razão social), percentual, valor absoluto e vencimento. (Estrutura exata não confirmada via MCP — ver PV-02.)"],
        ["Como o Deal é vinculado", "O ID do Negócio ou propriedade de controle é usado como chave de filtro na consulta HubDB. A lógica exata de filtragem está no módulo utils/divisao_faturamento."],
        ["Problema identificado", "Ao refazer a divisão de faturamento no Negócio, novas linhas são criadas no HubDB sem que as antigas sejam deletadas. Linhas com status 'em espera' causavam duplicação na minuta."],
        ["Solução em estabilização", "O time de contratos deletava manualmente as linhas com status 'em espera' no HubDB como paliativo. Correção definitiva foi solicitada (ticket 86ba2gkmc, done)."],
        ["Regras de conteúdo", "Para Upsell de usuários e sistemas: não incluir linhas de serviços (NFe, NFs, APIs) — apenas SaaS ou LU/MM/DC. Para Troca de API: apenas a linha da API. Para Upsell de serviços NFs: incluir linhas de serviços."],
    ],
)
set_col_widths(t45, [4, 13.5])
doc.add_paragraph()

# 4.6 Cards e Customizações
add_heading(doc, "4.6 Cards e Customizações", level=2)
add_para(doc,
    "A UI Extensions (card customizado no Negócio) é o ponto de entrada onde o vendedor ou CS preenche os dados "
    "de usuários, sistemas e divisão de faturamento antes de gerar a minuta. Os dados preenchidos na UI Extensions "
    "alimentam tanto os itens de linha do Negócio quanto a tabela HubDB de divisão de faturamento."
)
doc.add_paragraph()
add_para(doc,
    "Customização identificada: inclusão de coluna de vencimento para retrações de serviço (ticket 86ba46gup, blocked). "
    "Esta funcionalidade estava bloqueada no período de estabilização — verificar se foi implementada.",
    italic=True, size=10.5
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 5. RISCOS E DEPENDÊNCIAS
# ════════════════════════════════════════════════════════
add_heading(doc, "5. Riscos e Dependências", level=1)

t5 = create_brand_table(doc,
    ["Risco / Dependência", "Impacto", "Mitigação atual"],
    [
        ["Templates HubL não versionados formalmente",
         "Qualquer edição direta no Design Manager sem documentação altera o comportamento de todos os orçamentos futuros sem rastreabilidade.",
         "Processo depende de controle manual por Jorge Souza. Sem versionamento automático identificado."],
        ["Duplicação de linhas no HubDB",
         "Ao refazer a divisão de faturamento, linhas antigas não são deletadas. Causa duplicação na minuta.",
         "Correção foi solicitada (ticket 86ba2gkmc, done). Verificar se a solução definitiva foi implementada."],
        ["Coluna de vencimento em retrações de serviço ausente",
         "Minutas de retração de serviço não exibem a data de vencimento — informação exigida pelo time de contratos.",
         "Ticket 86ba46gup estava blocked. Verificar status atual."],
        ["Fluxo de geração automática de Quote desativado",
         "O workflow de geração automática de orçamento para retração (1793114638) está desabilitado. Se o CS esquecer de gerar o orçamento manualmente, o processo para.",
         "Processo operacional depende de disciplina do time de CS para selecionar o modelo e clicar em 'Assinar e Enviar'."],
        ["Propriedade 'Número do Contrato' legada",
         "Algumas minutas ainda puxavam de 'Número do Contrato' (com traço-aditivo) ao invés de 'Número do Contrato Gerado'. Tickets de estabilização corrigiram os casos identificados — mas podem existir outros.",
         "Revisar todos os templates que referenciam número de contrato para garantir uso da propriedade correta."],
        ["Signatários específicos por produto",
         "Minutas de CS para Sienge Plataforma requerem Maria Madalena e Giovani como signatários. Se essa configuração for alterada ou os usuários forem desativados no HubSpot, as minutas ficam incompletas.",
         "Configuração implementada via ticket 86b9ybfyr. Verificar se está parametrizado de forma dinâmica ou hardcoded."],
        ["Canal Starian — fluxo de aprovação específico",
         "Contratos de implantação do canal Starian não devem passar pelo aprovador financeiro. Solução definitiva foi solicitada (ticket 86b9ymc51, done) — verificar implementação.",
         "Havia contorno manual: remover o canal do Negócio para mudar o fluxo de aprovação."],
    ],
)
set_col_widths(t5, [5.5, 5, 7])
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 6. MATERIAIS DE APOIO
# ════════════════════════════════════════════════════════
add_heading(doc, "6. Materiais de Apoio", level=1)

t6mat = create_brand_table(doc,
    ["Tipo", "Título / ID", "Localização"],
    [
        ["Documento original", "_5.1 - Minutas [Dev].docx (32 páginas, mai/2026)", "docs/processos/5.1 - Minutas [Dev]/documentacao-gerada/"],
        ["Tickets ClickUp (principais)", "86bad86r6, 86bacxdtq, 86ba2gkmc, 86b9vy8kb, 86ba46gup", "ClickUp — projeto Sienge RaaS"],
        ["Tarefas filtradas", "5.1 - Minutas [Dev].md (54 tarefas)", "docs/processos/5.1 - Minutas [Dev]/clickup/"],
        ["Workflows HubSpot", "workflows-5.1.md", "docs/processos/5.1 - Minutas [Dev]/hubspot/"],
        ["Workflow [Minutas] 01", "Negócio Criado → Envia Infos da Empresa para Negócio", "HubSpot Portal 50102745 — id 1793711661"],
        ["Workflow [Minutas] 02", "Etapa = Contrato → Envia Infos do Negócio para Empresa", "HubSpot Portal 50102745 — id 1793738566"],
        ["Workflow fluxo ganho", "Nx. [Fluxo de envio para ganho manual] - Identificação do modelo", "HubSpot Portal 50102745 — id 1791307187"],
        ["Design Manager HubSpot", "Templates de orçamento personalizado", "HubSpot Portal 50102745 → Marketing → Arquivos e Modelos → Design Manager"],
        ["HubDB", "Tabela de divisão de faturamento", "HubSpot Portal 50102745 → Marketing → Conteúdo → HubDB"],
    ],
)
set_col_widths(t6mat, [3, 8, 6.5])
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 7. PONTOS A VALIDAR
# ════════════════════════════════════════════════════════
add_heading(doc, "7. Pontos a Validar", level=1)

add_callout(doc, "warning",
    "Ação requerida — Jorge Souza",
    "Os pontos abaixo não foram confirmados durante a documentação e requerem validação do responsável técnico "
    "(Jorge Souza) antes da publicação definitiva deste documento."
)
doc.add_paragraph()

t7 = create_brand_table(doc,
    ["ID", "Ponto", "Detalhe"],
    [
        ["PV-01", "Nomes exatos dos templates no Design Manager (CRÍTICO)",
         "Os 16 templates listados na seção 4.4 foram identificados nos tickets de estabilização e no DOCX original de mai/2026. "
         "Os nomes exatos como aparecem no Design Manager do HubSpot não foram confirmados via MCP (o MCP local não acessa o Design Manager). "
         "Requer que Jorge Souza liste os nomes exatos dos templates ativos no portal para atualizar a tabela 4.4."],
        ["PV-02", "Estrutura da tabela HubDB de divisão de faturamento",
         "A tabela HubDB que armazena a divisão de faturamento foi identificada pelos tickets de estabilização, mas sua estrutura "
         "(nome da tabela, colunas, chave de vínculo com o Negócio) não foi confirmada via MCP. "
         "Requer que Jorge Souza informe o nome da tabela e a lógica de filtro usada no template HubL."],
        ["PV-03", "Módulos HubL — nomes e estrutura de diretórios atuais",
         "A estrutura de módulos HubL listada na seção 4.4 (utils/, nova_venda/, upsell/, downsell/) foi documentada com base "
         "na versão original de mai/2026. Podem ter ocorrido renomeações, adições ou remoções de módulos durante a estabilização. "
         "Verificar com Jorge Souza se a estrutura atual do Design Manager corresponde ao que está documentado."],
        ["PV-04", "Status da coluna de vencimento para retrações de serviço (ticket 86ba46gup)",
         "Este ticket estava com status 'blocked' — a coluna de vencimento na UI Extensions para retrações de serviço "
         "não havia sido implementada. Verificar se foi desbloqueada e implementada após o período de estabilização."],
        ["PV-05", "Solução definitiva para duplicação de linhas no HubDB (ticket 86ba2gkmc)",
         "Ticket marcado como 'done', mas a solução paliativa era deletar linhas manualmente. "
         "Confirmar com Jorge Souza se a correção definitiva (prevenção da duplicação ao refazer a divisão de faturamento) "
         "foi implementada no código do template ou em um workflow de limpeza."],
        ["PV-06", "Signatários Maria Madalena e Giovani — configuração dinâmica ou hardcoded",
         "O ticket 86b9ybfyr (done) adicionou esses signatários às minutas de CS para Sienge Plataforma. "
         "Confirmar se a configuração é dinâmica (lida de uma propriedade ou lista no HubSpot) ou hardcoded no template HubL, "
         "para avaliar o risco de quebra caso os usuários sejam desativados."],
        ["PV-07", "Fluxo de aprovação canal Starian — implementação definitiva (ticket 86b9ymc51)",
         "O ticket foi resolvido como 'done', mas havia contorno manual em andamento. "
         "Confirmar com Jorge Souza ou Vinicius Vieira Braz se a regra de exclusão do time financeiro nas aprovações "
         "de implantação Starian (e de não reprovação automática do Gestor Obras) foi implementada definitivamente no "
         "workflow 1790676277 ou em outro fluxo."],
        ["PV-08", "Workflow de geração automática de Quote (1793114638) — substituição definitiva",
         "O workflow está marcado como 'desatualizado' e desabilitado. Confirmar se o fluxo manual atual é definitivo "
         "ou se existe plano de reimplementar a geração automática de orçamento para retração."],
        ["PV-09", "Propriedades lidas no HubL — Deal vs. Empresa",
         "Os workflows [Minutas] 01 e 02 copiam propriedades entre Empresa e Negócio. "
         "Não foi confirmado quais propriedades são lidas diretamente do objeto Empresa no HubL "
         "(via associated_objects) versus quais são lidas do Negócio após a cópia. "
         "Impacta o diagnóstico de minutas com dados incorretos (ex: CNPJ errado na minuta)."],
    ],
)
set_col_widths(t7, [1.5, 4.5, 11.5])
doc.add_paragraph()

doc.save(OUTPUT_PATH)
print(f"Documento gerado: {OUTPUT_PATH}")
