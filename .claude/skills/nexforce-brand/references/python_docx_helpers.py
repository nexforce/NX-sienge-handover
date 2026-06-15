"""
Nexforce brand helpers for python-docx.

Usage:
    from docx import Document
    from python_docx_helpers import (
        apply_nexforce_theme,
        insert_brand_header,
        insert_brand_footer,
        insert_brand_cover,
        create_brand_table,
        apply_semantic_text,
        add_callout,
        NEXFORCE_COLORS,
        logo_path,
    )

    doc = Document()
    apply_nexforce_theme(doc)
    insert_brand_header(doc, title="Q1 Results", period="Jan-Mar 2026")
    insert_brand_footer(doc)

    create_brand_table(
        doc,
        headers=["Metric", "Target", "Actual", "Delta"],
        rows=[
            ["Revenue", "1.2M", "1.38M", "+15%"],
            ["Churn",   "3.0%", "2.4%",  "-0.6pp"],
        ],
        total_row=["Total", "1.2M", "1.38M", "+15%"],
    )
    doc.save("report.docx")

All functions preserve the Nexforce brand rules: Lato font, palette-only colors,
semantic color never on cell backgrounds, and logos resolved from the skill's
local assets folder.
"""

from __future__ import annotations

import glob
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.document import Document as DocType
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell, Table


# === Palette ==============================================================

NEXFORCE_COLORS: dict[str, str] = {
    "primary_black": "0C0E0E",
    "white":         "FFFFFF",
    "dark_gray":     "515151",
    "mid_gray":      "777777",
    "light_gray":    "9C9B9B",
    "near_white":    "F5F5F5",
    "navy":          "303F63",
    "blue":          "215A9F",
    "green":         "2D6E44",
    "yellow":        "D8B523",
    "red":           "BA1925",
    "orange":        "D56316",
    "purple":        "662A7F",
    "pink":          "B1338A",
    # Callout backgrounds (tinted from accent colors)
    "callout_info_bg":      "EEF3FB",
    "callout_success_bg":   "EEF7F1",
    "callout_warning_bg":   "FDF8E1",
    "callout_alert_bg":     "FDECEA",
    "callout_strategic_bg": "ECEEF4",
}


def _rgb(key: str) -> RGBColor:
    return RGBColor.from_string(NEXFORCE_COLORS[key])


# === Asset paths ==========================================================

def nexforce_assets() -> Path:
    """Locate the nexforce-brand/assets folder across environments."""
    candidates = [
        # 1) claude.ai computer-use (bash_tool sandbox)
        Path("/mnt/skills/user/nexforce-brand/assets"),
        # 2) Cowork Linux sandbox
        *[Path(p) for p in glob.glob("/sessions/*/mnt/Work/Skills/nexforce-brand/assets")],
        *[Path(p) for p in glob.glob("/sessions/*/mnt/.claude/skills/nexforce-brand/assets")],
        # 3) Windows local
        Path.home() / "Desktop" / "Work" / "Skills" / "nexforce-brand" / "assets",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(
        "nexforce-brand/assets folder not found. Check skill installation."
    )


def logo_path(variant: str) -> str:
    """Absolute path to logo-black.png or logo-white.png."""
    if variant not in ("black", "white"):
        raise ValueError("variant must be 'black' or 'white'")
    return str(nexforce_assets() / f"logo-{variant}.png")


# === Theme ================================================================

def apply_nexforce_theme(doc: DocType) -> None:
    """Set Lato as the theme font and apply palette colors to default styles.

    Falls back to Calibri if Lato is not available at render time.
    Safe to call on a fresh Document().
    """
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Lato"
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = _rgb("dark_gray")

    # East-Asia / complex-script fallback so Lato is set for all scripts
    rpr = style_normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), "Lato")

    heading_map = {
        "Heading 1": ("primary_black", 24, True),
        "Heading 2": ("navy",          18, True),
        "Heading 3": ("dark_gray",     14, True),
        "Title":     ("primary_black", 32, True),
    }
    for name, (color_key, size, bold) in heading_map.items():
        try:
            s = doc.styles[name]
        except KeyError:
            continue
        s.font.name = "Lato"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = _rgb(color_key)


# === Header / Footer ======================================================

def insert_brand_header(doc: DocType, title: str,
                         period: Optional[str] = None) -> None:
    """Top-left black logo with optional right-aligned document info."""
    section = doc.sections[0]
    header = section.header

    p = header.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(logo_path("black"), height=Cm(1.2))

    if title:
        info = title + (f"   ·   {period}" if period else "")
        p.add_run("\t\t")
        text_run = p.add_run(info)
        text_run.font.name = "Lato"
        text_run.font.size = Pt(9)
        text_run.font.color.rgb = _rgb("mid_gray")


def insert_brand_footer(doc: DocType,
                         confidentiality: str = "Uso interno") -> None:
    """Text-only footer, right-aligned, color #9C9B9B. No logo per brand rules."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run(confidentiality)
    run.font.name = "Lato"
    run.font.size = Pt(8)
    run.font.color.rgb = _rgb("light_gray")


# === Cover page ===========================================================

def insert_brand_cover(doc: DocType, title: str,
                        subtitle: Optional[str] = None,
                        author: Optional[str] = None) -> None:
    """Insert a light-background cover page with the black logo and large title.

    Note: python-docx cannot reliably recolor the Word page background across
    versions. For a true dark cover, render the cover as an image or generate
    cover-heavy decks as .pptx instead. This implementation produces a
    polished light cover that is safe on any Word installation.
    """
    first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    first.paragraph_format.space_before = Cm(6)

    run = first.add_run()
    run.add_picture(logo_path("black"), height=Cm(2))

    doc.add_paragraph()  # spacer

    t = doc.add_paragraph()
    t.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = t.add_run(title)
    tr.font.name = "Lato"
    tr.font.size = Pt(32)
    tr.font.bold = True
    tr.font.color.rgb = _rgb("primary_black")

    if subtitle:
        s = doc.add_paragraph()
        sr = s.add_run(subtitle)
        sr.font.name = "Lato"
        sr.font.size = Pt(14)
        sr.font.color.rgb = _rgb("dark_gray")

    if author:
        a = doc.add_paragraph()
        a.paragraph_format.space_before = Cm(1)
        ar = a.add_run(author)
        ar.font.name = "Lato"
        ar.font.size = Pt(10)
        ar.font.color.rgb = _rgb("mid_gray")

    doc.add_page_break()


# === Tables ===============================================================

def _set_cell_bg(cell: _Cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _style_cell_text(cell: _Cell, *, bold: bool = False,
                      color_key: str = "dark_gray",
                      size: int = 10) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "Lato"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = _rgb(color_key)


def create_brand_table(doc: DocType, headers: List[str],
                        rows: List[List[str]],
                        total_row: Optional[List[str]] = None) -> Table:
    """Branded table: #0C0E0E header row, alternating #FFFFFF/#F5F5F5 body,
    optional total row on #0C0E0E. Never fills data cells with semantic color —
    use apply_semantic_text() on a specific cell after creation for that.
    """
    n_cols = len(headers)
    n_rows = 1 + len(rows) + (1 if total_row else 0)
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # Header row
    _set_row_cant_split(table.rows[0])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        _set_cell_bg(cell, NEXFORCE_COLORS["primary_black"])
        _style_cell_text(cell, bold=True, color_key="white", size=10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_values in enumerate(rows):
        bg_key = "white" if r_idx % 2 == 0 else "near_white"
        _set_row_cant_split(table.rows[r_idx + 1])
        for c_idx, value in enumerate(row_values):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            _set_cell_bg(cell, NEXFORCE_COLORS[bg_key])
            _style_cell_text(cell, color_key="dark_gray", size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Total row
    if total_row:
        last_idx = n_rows - 1
        _set_row_cant_split(table.rows[last_idx])
        for c_idx, value in enumerate(total_row):
            cell = table.rows[last_idx].cells[c_idx]
            cell.text = str(value)
            _set_cell_bg(cell, NEXFORCE_COLORS["primary_black"])
            _style_cell_text(cell, bold=True, color_key="white", size=10)

    return table


def _set_row_cant_split(row) -> None:
    """Mark a table row as 'do not split across pages' (Word w:cantSplit).

    Mirrors the CSS `break-inside: avoid` rule documented in SKILL.md
    § Pagination & Layout Safety.
    """
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def _set_paragraph_keep_with_next(paragraph) -> None:
    """Mark a paragraph so Word never splits it from the next block.

    Mirrors the CSS `break-after: avoid` rule on headings.
    """
    pPr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    pPr.append(keep_next)
    keep_lines = OxmlElement("w:keepLines")
    pPr.append(keep_lines)


def apply_semantic_text(cell: _Cell, value: str, kind: str) -> None:
    """Color the text inside a data cell based on semantic meaning.

    Never touches the cell background — brand rule requires white/near-white
    body cells, color on text only.

    kind ∈ {positive, attention, alert, info}
        positive  -> green  #2D6E44 (AA on white)
        attention -> yellow #D8B523 (AVOID on light; legible only on dark)
        alert     -> red    #BA1925 (AA on white)
        info      -> blue   #215A9F (AA on white)
    """
    mapping = {
        "positive":  "green",
        "attention": "yellow",
        "alert":     "red",
        "info":      "blue",
    }
    if kind not in mapping:
        raise ValueError(
            f"Unknown kind {kind!r}; allowed: {sorted(mapping.keys())}"
        )

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(value)
    run.font.name = "Lato"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = _rgb(mapping[kind])


# === Callouts =============================================================

_CALLOUT_CONFIG = {
    "info":      ("blue",   "callout_info_bg"),
    "success":   ("green",  "callout_success_bg"),
    "warning":   ("yellow", "callout_warning_bg"),
    "alert":     ("red",    "callout_alert_bg"),
    "strategic": ("navy",   "callout_strategic_bg"),
}


def add_callout(doc: DocType, kind: str, title: str, body: str) -> None:
    """Single-cell table styled as a left-border callout block.

    kind ∈ {info, success, warning, alert, strategic}
    """
    if kind not in _CALLOUT_CONFIG:
        raise ValueError(
            f"Unknown kind {kind!r}; allowed: {sorted(_CALLOUT_CONFIG.keys())}"
        )
    border_key, bg_key = _CALLOUT_CONFIG[kind]

    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _set_cell_bg(cell, NEXFORCE_COLORS[bg_key])

    # Pagination safety — callout never splits across pages.
    # See SKILL.md § Pagination & Layout Safety.
    _set_row_cant_split(t.rows[0])

    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")  # 24 eighths-of-a-pt ≈ 3pt thick border
    left.set(qn("w:color"), NEXFORCE_COLORS[border_key])
    borders.append(left)
    tcPr.append(borders)

    cell.text = ""
    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    r.font.name = "Lato"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = _rgb(border_key)
    # Title sticks to body — never orphaned at the bottom of a page.
    _set_paragraph_keep_with_next(p_title)

    p_body = cell.add_paragraph()
    r = p_body.add_run(body)
    r.font.name = "Lato"
    r.font.size = Pt(10)
    r.font.color.rgb = _rgb("dark_gray")
